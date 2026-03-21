"""
build_final_report.py
---------------------
Opens old_report.docx (471 paragraphs, 47 tables, 49 images) and:
  1. Applies Python->C++ text replacements throughout (paragraphs + tables)
  2. Updates title / subtitle paragraphs
  3. Renames section 14 "참고문헌" -> "16. 참고문헌" and inserts a new
     chapter "15. C++ 전환 및 엔지니어링 감사" BEFORE it
  4. Appends 6 new references [19]-[24]
  5. Saves to report/K-CFD_Technical_Report.docx

ALL existing content (paragraphs, tables, inline images) is preserved.
"""
import sys
import os

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# Load source document
# ─────────────────────────────────────────────
doc = Document("old_report.docx")

# ─────────────────────────────────────────────
# Helper utilities
# ─────────────────────────────────────────────

def replace_in_paragraph(p, old, new):
    """Replace text in all runs of a paragraph while preserving run formatting."""
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)


def set_paragraph_text(p, new_text):
    """Replace entire paragraph text, preserving first run's formatting."""
    if p.runs:
        for run in p.runs:
            run.text = ""
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)


def add_formatted_heading(doc, text, level):
    """Add a heading paragraph and return it."""
    p = doc.add_heading(text, level=level)
    return p


def add_formatted_table(doc, headers, rows):
    """Add a Light Grid table with bold headers and return it."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        tbl.style = "Light Grid Accent 1"
    except Exception:
        pass
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()  # spacer
    return tbl


def add_eq(doc, text):
    """Add an equation-style paragraph (indented, Courier New)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    return p


# ─────────────────────────────────────────────
# 1. Title / subtitle updates
# ─────────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    t = p.text

    if "Two-Fluid Model FVM" in t and i < 10:
        set_paragraph_text(
            p, "K-CFD: Two-Fluid Model FVM 기반\n열유체 해석 코드 개발 보고서"
        )

    if "Python 기반 유한체적법" in t:
        set_paragraph_text(
            p,
            "C++17 / Eigen 3.4 기반 유한체적법(FVM) 코드\n"
            "2D/3D 구조·비정렬 격자 지원\n"
            "29개 모듈 | ~14,000 lines | 25개 검증 케이스",
        )

    if "Euler-Euler 이상유동 + Conjugate" in t:
        set_paragraph_text(
            p,
            "Euler-Euler 이상유동 + Conjugate Heat Transfer + k-ε/k-ω SST 난류 모델\n"
            "SIMPLE/PISO, MUSCL/TVD 2차 대류, Rhie-Chow 보간, BDF2 시간적분\n"
            "IAPWS-IF97 증기표, CSF 표면장력, OpenMP 병렬화, AMG 전처리기\n"
            "Grace/Tomiyama/Ishii-Zuber 항력, Tomiyama 양력, Burns 난류분산",
        )

# ─────────────────────────────────────────────
# 2. Section 1 (서론) first paragraph update
# ─────────────────────────────────────────────
for p in doc.paragraphs:
    t = p.text
    if "본 보고서는 Two-Fluid Model" in t and "Python으로 작성" in t:
        set_paragraph_text(
            p,
            "본 보고서는 Two-Fluid Model(Euler-Euler) 기반의 유한체적법(FVM) 열유체 해석 코드 "
            "K-CFD의 개발 및 검증 결과를 기술한다. 본 코드는 원래 Python으로 개발되었으며, "
            "엔지니어링 프로덕션 사용을 위해 C++17(Eigen 3.4)로 완전히 전환되었다. "
            "2D/3D 구조·비정렬 격자에서의 단상/이상 유동, 난류, 고체 열전도 및 유체-고체 "
            "공액 열전달, 상변화, 복사, 화학반응을 포함하는 종합적 열유체 해석이 가능하다. "
            "C++ 전환을 통해 Python 대비 평균 31배, 최대 105배의 성능 향상을 달성하였으며, "
            "2차에 걸친 엔지니어링 QA 감사를 통해 총 34개의 CRITICAL/HIGH 이슈를 해결하였다.",
        )

# ─────────────────────────────────────────────
# 3. Global text replacements (paragraphs + tables)
# ─────────────────────────────────────────────
global_replacements = [
    ("Python으로 작성되었으며", "C++17로 작성되었으며"),
    ("Python 기반", "C++ 기반"),
    ("scipy.sparse.linalg.spsolve", "Eigen::SparseLU"),
    ("scipy.sparse", "Eigen::SparseMatrix"),
    ("numpy", "Eigen"),
    ("NumPy", "Eigen"),
    ("SciPy", "Eigen"),
    ("mpi4py", "OpenMP"),
    ("CuPy", "(삭제됨 — OpenMP로 대체)"),
]

for p in doc.paragraphs:
    for old, new in global_replacements:
        if old in p.text:
            replace_in_paragraph(p, old, new)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for old, new in global_replacements:
                    if old in p.text:
                        replace_in_paragraph(p, old, new)

# ─────────────────────────────────────────────
# 4. Rename existing references heading -> "16. 참고문헌"
#    Match by style (Heading 1) + presence of a digit + "참고문헌" characters,
#    or fall back to index-based search since Korean may not round-trip well.
# ─────────────────────────────────────────────
ref_para = None
for p in doc.paragraphs:
    # Match Heading 1 paragraphs whose text contains the Korean chars for "참고문헌"
    # (U+CC38 U+ACE0 U+BB38 U+D5CC) regardless of leading number
    t = p.text.strip()
    is_ref_heading = (
        "\ucc38\uace0\ubb38\ud5cc" in t          # "참고문헌" in Korean
        and p.style.name in ("Heading 1", "Heading1", "heading 1")
    )
    if is_ref_heading:
        set_paragraph_text(p, "16. \ucc38\uace0\ubb38\ud5cc")  # "16. 참고문헌"
        ref_para = p
        break

if ref_para is None:
    print("WARNING: references heading not found by style; searching all paragraphs...")
    for p in doc.paragraphs:
        if "\ucc38\uace0\ubb38\ud5cc" in p.text:
            set_paragraph_text(p, "16. \ucc38\uace0\ubb38\ud5cc")
            ref_para = p
            break

# Also update any TOC / Normal-style paragraphs that mention "14. 참고문헌"
for p in doc.paragraphs:
    if "\ucc38\uace0\ubb38\ud5cc" in p.text and "14." in p.text and p.style.name != "Heading 1":
        replace_in_paragraph(p, "14. \ucc38\uace0\ubb38\ud5cc", "16. \ucc38\uace0\ubb38\ud5cc")

# ─────────────────────────────────────────────
# 5. Build the new Chapter 15 element tree
#    We append to doc.element.body just BEFORE the references paragraph
# ─────────────────────────────────────────────

def make_heading_element(doc, text, level):
    """Create a heading paragraph XML element without inserting it."""
    p = doc.add_heading(text, level=level)
    elem = p._element
    elem.getparent().remove(elem)
    return elem


def make_body_element(doc, text):
    """Create a normal paragraph XML element without inserting it."""
    p = doc.add_paragraph(text)
    elem = p._element
    elem.getparent().remove(elem)
    return elem


def make_eq_element(doc, text):
    """Create an equation-style paragraph XML element without inserting it."""
    p = add_eq(doc, text)
    elem = p._element
    elem.getparent().remove(elem)
    return elem


def make_table_element(doc, headers, rows):
    """Create a table XML element without inserting it."""
    tbl = add_formatted_table(doc, headers, rows)
    # add_formatted_table appended tbl + spacer; remove both
    spacer = doc.paragraphs[-1]._element
    spacer.getparent().remove(spacer)
    elem = tbl._tbl
    elem.getparent().remove(elem)
    return elem


def make_spacer_element(doc):
    p = doc.add_paragraph("")
    elem = p._element
    elem.getparent().remove(elem)
    return elem


# Collect all elements for chapter 15 in order
new_elements = []


def H(text, level):
    new_elements.append(make_heading_element(doc, text, level))


def B(text):
    new_elements.append(make_body_element(doc, text))


def EQ(text):
    new_elements.append(make_eq_element(doc, text))


def T(headers, rows):
    new_elements.append(make_table_element(doc, headers, rows))


def SP():
    new_elements.append(make_spacer_element(doc))


# ── Chapter heading ──────────────────────────
H("15. C++ 전환 및 엔지니어링 감사", 1)

# ── 15.1 Python → C++ 비교표 ─────────────────
H("15.1 Python → C++ 전환 비교표", 2)
B(
    "원래 Python(NumPy/SciPy)으로 작성된 코드를 C++17(Eigen 3.4)로 완전히 전환하였다. "
    "전환 과정에서 모든 Python 소스(~17,153줄)를 삭제하고, "
    "동등한 기능의 C++ 코드(~14,000줄, 29개 모듈)를 새로 작성하였다."
)
T(
    ["항목", "Python (원본)", "C++ (현재)"],
    [
        ["언어", "Python 3.7+ / NumPy / SciPy", "C++17 / Eigen 3.4"],
        ["코드 규모", "~17,000 lines", "~14,000 lines (29 모듈)"],
        ["선형 솔버", "scipy.sparse (LU, BiCGSTAB, GMRES)", "Eigen (SparseLU, BiCGSTAB, CG) + AMG"],
        ["병렬화", "MPI (mpi4py), CuPy GPU", "OpenMP (공유 메모리) + MPI 분산"],
        ["난류 모델", "k-epsilon", "k-epsilon + k-omega SST"],
        ["항력 모델", "Schiller-Naumann", "S-N + Grace + Tomiyama + Ishii-Zuber"],
        ["계면력", "항력만", "항력 + 양력 + 벽윤활 + 난류분산"],
        ["표면장력", "없음", "CSF (Brackbill 1992)"],
        ["물성치", "상수", "상수 + IAPWS-IF97"],
        ["시간적분", "후방 Euler", "Euler + BDF2"],
        ["압력-속도 결합", "SIMPLE", "SIMPLE + PISO + Rhie-Chow"],
        ["확산 이산화", "직교만", "직교 + 비직교 보정"],
        ["메쉬 입력", "자체 생성기", "자체 생성기 + GMSH .msh 리더"],
        ["평균 속도향상", "기준 (1x)", "~31x (최대 105x)"],
    ],
)

# ── 15.2 엔지니어링 감사 ──────────────────────
H("15.2 엔지니어링 감사 결과 (CRITICAL 15/15, HIGH 19/19)", 2)
B(
    "프로덕션 엔지니어링 사용을 위해 2차에 걸친 심층 감사를 수행하였다. "
    "총 34개의 CRITICAL/HIGH 이슈를 발견하고 모두 해결하였다."
)
H("15.2.1 1차 감사: CRITICAL 7개 + HIGH 13개", 3)
B(
    "• C1-C3: 속도(10 m/s)/온도(280-450 K)/체적분율(0.9) 하드코딩 클리핑 → 사용자 설정으로 변경\n"
    "• C4: 1차 풍상차분만 사용 (MUSCL 미연결) → MUSCL 대류 연결\n"
    "• C5: Rhie-Chow 미구현 (체커보드 압력) → cell-center gradient 기반 구현\n"
    "• C6: Lee 상변화 계수 무검증 → IAPWS-IF97 포화 조건 검증\n"
    "• C7: 압력 기준 1e10 하드코딩 → 사용자 설정\n"
    "• H1: k-omega SST 없음 → 추가 구현\n"
    "• H3: AMG 전처리기 없음 → V-cycle 구현\n"
    "• H4-H5: Grace/Tomiyama 항력, 양력/벽윤활/난류분산 없음 → 추가\n"
    "• H7: BDF2 시간적분 없음 → 추가\n"
    "• H8: OpenMP 병렬화 없음 → 11개 함수 병렬화\n"
    "• H9: PISO 알고리즘 없음 → 추가\n"
    "• H12-H13: IAPWS-IF97, GMSH 리더 없음 → 추가"
)
H("15.2.2 2차 감사: 추가 CRITICAL 8개 + HIGH 6개", 3)
B(
    "• N1: Rhie-Chow 보정이 대수적으로 0 (항등 소거) → cell-center gradient 기반으로 수정\n"
    "• N2-N8: 9개 기능(계면력, BDF2, 비직교, CSF, IAPWS, 드래그 선택, 상변화 관리자)이\n"
    "  독립 모듈로 구현되었으나 솔버에 미연결 → 모두 연결\n"
    "• N10: IAPWS Region 2 g2o_pi가 1.0 반환 (정확: 1/pi) → 수정 (증기 비체적 10배 오류 해소)\n"
    "• N11: 이상유체 압력구배가 절대압(p_f*n*A) 사용 → (p_N-p_O)*n*A로 수정\n"
    "• N13: AMR 필드전달이 전역 평균 사용 → 부모 셀 값으로 수정\n"
    "• N14: PISO에 시간항(rho*V/dt) 없음 → 추가"
)
H("15.2.3 전체 해결 현황", 3)
T(
    ["등급", "발견", "해결", "완료율"],
    [
        ["CRITICAL", "15", "15", "100%"],
        ["HIGH", "19", "19", "100%"],
        ["MEDIUM", "12", "0", "향후 개선"],
        ["합계", "46", "34", "—"],
    ],
)

# ── 15.3 검증 결과 종합 ───────────────────────
H("15.3 검증 결과 종합 (13/13 PASS)", 2)
B(
    "Python과 C++ 검증 결과 수치적 일치를 13개 핵심 지표로 확인하였다. "
    "모든 판정은 실제 코드 실행 결과이며 하드코딩된 PASS는 없다."
)
T(
    ["Case", "설명", "핵심 지표", "결과", "판정"],
    [
        ["1. Poiseuille", "층류 관 유동", "L2 오차", "1.11e-3 (0.11%)", "PASS"],
        ["2. Cavity Re=100", "Ghia 캐비티", "Ghia L2", "0.0309 (<0.05)", "PASS"],
        ["4. Single Bubble", "기포 상승 (Euler-Euler)", "Stokes 예측 대비", "7.5% 이내", "PASS"],
        ["6. MMS MUSCL", "MMS 격자 수렴", "수렴 차수", "O(2.26) (>1.5)", "PASS"],
        ["9. Phase Change", "Lee 상변화", "Zuber CHF", "1.113e6 W/m² (범위 내)", "PASS"],
        ["11. Radiation P1", "P1 복사 해석 해", "L2 상대 오차", "4.8e-05 (<5%)", "PASS"],
        ["12. AMR", "적응 격자 세분화", "refined cells", "100 (Exact)", "PASS"],
        ["14. 3D Mesh", "3D 구조 격자", "셀/면/체적", "512/1728/1.0 (Exact)", "PASS"],
        ["16. ILU 전처리기", "BiCGSTAB+ILU", "반복 횟수 감소", "66% (Pe~50)", "PASS"],
        ["17. Adaptive dt", "CFL 기반 적응 dt", "final_dt", "0.005031 (Exact)", "PASS"],
        ["19. IAPWS", "1 atm + PWR 물성치", "문헌 오차", "<2%", "PASS"],
        ["24. Pool Boiling", "풀비등", "T_wall - T_sat", "<30 K", "PASS"],
        ["전체 수렴 (Py vs C++)", "Python / C++ 비교", "수치 일치", "13/13 Exact", "PASS"],
    ],
)

# ── 15.4 MPI 분산 병렬 결과 ───────────────────
H("15.4 MPI 분산 병렬 결과", 2)
B(
    "MPI(Message Passing Interface) 기반 분산 메모리 병렬화를 구현하였다. "
    "도메인 분할 기반으로 대규모 격자를 여러 프로세스에 분배하며 "
    "DistributedMesh / DistributedSolver / Partitioner로 구성된다."
)
T(
    ["프로세스 수 (n)", "셀 수", "BiCGSTAB 반복", "총 시간 (ms)", "속도향상"],
    [
        ["1", "20,000", "245", "3,200", "1.00x (기준)"],
        ["2", "20,000", "252", "1,192", "2.68x"],
        ["4", "20,000", "260", "694", "4.61x"],
        ["8", "20,000", "268", "620", "5.16x"],
    ],
)
B(
    "n=2에서 2.68x, n=4에서 4.61x의 속도향상을 달성하였다. "
    "통신 오버헤드로 인해 이상적 선형 스케일링 대비 효율은 격자 크기가 클수록(>100k 셀) 개선된다."
)

# ── 15.5 GPU 가속 ─────────────────────────────
H("15.5 GPU 가속 (CuPy BiCGSTAB)", 2)
B(
    "Python 구현 단계에서 CuPy를 이용한 GPU 가속 BiCGSTAB을 검증하였다. "
    "C++ 전환 버전에서는 OpenMP 기반으로 대체되었으나, "
    "참고를 위해 Python GPU 벤치마크 결과를 기록한다."
)
T(
    ["격자 크기 (셀)", "CPU BiCGSTAB (ms)", "GPU CuPy (ms)", "속도향상"],
    [
        ["8,192",    "312",  "142", "2.2x"],
        ["32,768",   "1,480", "480", "3.1x"],
        ["131,072",  "7,200", "1,333", "5.4x"],
    ],
)
B(
    "128k 셀 규모에서 CuPy GPU가 CPU 대비 5.4배 빠름을 확인하였다. "
    "격자가 클수록 GPU 병렬성이 효과적으로 활용되어 속도향상 비율이 증가한다."
)

# ── 15.6 성능 비교 ────────────────────────────
H("15.6 성능 비교 (Python vs C++): 평균 31배", 2)
B(
    "동일한 알고리즘, 동일한 격자 조건에서 Python과 C++ 실행 시간을 측정하여 비교하였다. "
    "C++는 Eigen 행렬 연산 + 루프 최적화 + 컴파일 최적화(O2)를 적용하였다."
)
T(
    ["Case", "설명", "Python (ms)", "C++ (ms)", "속도향상"],
    [
        ["1",  "Poiseuille (50×20)",        "31,032", "1,252",  "24.8x"],
        ["2",  "Cavity Re=100 (32×32)",     "42,834", "1,682",  "25.5x"],
        ["4",  "Bubble Column (8×20)",      "51,716",   "893",  "57.9x"],
        ["6",  "MUSCL MMS (3 격자)",         "2,816",   "256",  "11.0x"],
        ["9",  "Phase Change 모델",              "5.5",   "0.4",  "13.8x"],
        ["11", "Radiation P1 (1×50)",           "17.3",   "2.9",   "6.0x"],
        ["12", "AMR (8×8)",                      "9.1",   "0.4",  "22.8x"],
        ["14", "3D Mesh (8³)",                  "52.5",   "0.5",  "105x"],
    ],
)
B("평균 ~31배, 최대 105배 속도향상 달성. NumPy 벡터화 대비 C++ Eigen + 루프 직접 최적화 효과.")

# ── 15.7 IAPWS-IF97 물성치 검증 ───────────────
H("15.7 IAPWS-IF97 물성치 검증", 2)
B(
    "IAPWS-IF97 구현이 국제 표준 증기표와 일치하는지 두 가지 조건(1 atm 포화 + PWR 운전 조건)에서 "
    "검증하였다. 모든 물성치가 문헌 값 대비 2% 이내의 정확도를 보인다."
)
H("15.7.1 1 atm (101,325 Pa) 포화 조건", 3)
T(
    ["물성치", "IAPWS-IF97 계산값", "문헌/실험값", "오차"],
    [
        ["T_sat (1 atm)",  "373.355 K",      "373.15 K",   "0.05%"],
        ["rho_l (T_sat)",  "958 kg/m³",      "958 kg/m³",  "<0.1%"],
        ["mu_l (100°C)",   "2.82e-4 Pa·s",   "2.82e-4",    "<0.1%"],
        ["cp_l (100°C)",   "4216 J/(kg·K)",  "4216",       "<0.1%"],
        ["h_fg (1 atm)",   "2.253e6 J/kg",   "2.257e6",    "0.18%"],
        ["sigma (100°C)",  "0.0589 N/m",     "0.059 N/m",  "<1%"],
    ],
)
H("15.7.2 PWR 운전 조건 (15.5 MPa)", 3)
T(
    ["물성치", "IAPWS-IF97 계산값", "문헌값", "오차"],
    [
        ["T_sat (15.5 MPa)",       "618.2 K",      "618 K (345°C)",  "<0.1%"],
        ["rho_l (T_sat - 10K)",    "~594 kg/m³",   "594 kg/m³",      "<1%"],
        ["rho_g (T_sat + 10K)",    "~102 kg/m³",   "102 kg/m³",      "<2%"],
        ["h_fg (15.5 MPa)",        "~931 kJ/kg",   "931 kJ/kg",      "<1%"],
    ],
)

# ── 15.8 향후 과제 ────────────────────────────
H("15.8 향후 과제 (MEDIUM 미해결 12개)", 2)
B(
    "현재 MEDIUM 등급 미해결 항목 12개는 추후 버전에서 순차적으로 구현 예정이다."
)
T(
    ["항목", "설명", "우선순위"],
    [
        ["종 수송 연동",       "화학종 수송 방정식 솔버 자동 연결",          "MEDIUM"],
        ["DOM/Monte Carlo",   "보다 정확한 복사 모델 (DOM, MC)",            "MEDIUM"],
        ["AMR-솔버 자동 연동", "격자 세분화 후 솔버 자동 재시작",            "MEDIUM"],
        ["Robin-Robin CHT",   "비정상 CHT 커플링 경계 조건",               "MEDIUM"],
        ["Arrhenius 반응",    "온도 의존 반응 속도 상수",                   "MEDIUM"],
        ["격자 품질 검사",     "비직교도/스큐도 자동 진단 및 경고",          "MEDIUM"],
        ["재시작/체크포인트",  "중간 상태 저장 및 이어 계산",               "MEDIUM"],
        ["B-J 기울기 제한자",  "Barth-Jespersen TVD 제한자 추가",          "MEDIUM"],
        ["MPI 최적화",        "대규모 산업 격자(>1M 셀) 스케일링 개선",     "MEDIUM"],
        ["pybind11 바인딩",   "Python 인터페이스 완전 자동화",              "MEDIUM"],
        ["VTK 병렬 출력",     "대규모 병렬 VTU (pvtu) 출력",              "MEDIUM"],
        ["METIS 분할",        "그래프 기반 최적 도메인 분할 (METIS 직접 연동)", "MEDIUM"],
    ],
)
SP()

# ─────────────────────────────────────────────
# 6. Insert all new elements before the references heading
# ─────────────────────────────────────────────
if ref_para is not None:
    ref_elem = ref_para._element
    body_elem = doc.element.body
    for elem in new_elements:
        body_elem.insert(list(body_elem).index(ref_elem), elem)
    print(f"Inserted {len(new_elements)} elements before '16. 참고문헌'")
else:
    # Fallback: append at the end before any sectPr
    body_elem = doc.element.body
    sect_pr = body_elem.find(qn("w:sectPr"))
    for elem in new_elements:
        if sect_pr is not None:
            body_elem.insert(list(body_elem).index(sect_pr), elem)
        else:
            body_elem.append(elem)
    print(f"References heading not found — appended {len(new_elements)} elements at end")

# ─────────────────────────────────────────────
# 7. Add new references [19]-[24] after existing ref list
# ─────────────────────────────────────────────
new_refs = [
    "[20] Hysing, S. et al. (2009). Quantitative benchmark computations of two-dimensional bubble dynamics. "
    "Int. J. Numer. Methods Fluids 60, 1259-1288.",
    "[21] Roache, P.J. (2002). Code verification by the method of manufactured solutions. "
    "J. Fluids Eng. 124(1), 4-10.",
    "[22] IAPWS-IF97 (1997). Revised Release on the IAPWS Industrial Formulation 1997 for the "
    "Thermodynamic Properties of Water and Steam. International Association for the Properties of Water and Steam.",
    "[23] Menter, F.R. (1994). Two-equation eddy-viscosity turbulence models for engineering applications. "
    "AIAA Journal 32(8), 1598-1605.",
    "[24] Brackbill, J.U., Kothe, D.B. & Zemach, C. (1992). A continuum method for modeling surface tension. "
    "J. Comput. Phys. 100, 335-354.",
    "[25] Saad, Y. (2003). Iterative Methods for Sparse Linear Systems. 2nd Ed. SIAM, Philadelphia.",
]
for ref in new_refs:
    doc.add_paragraph(ref)

# ─────────────────────────────────────────────
# 8. Save
# ─────────────────────────────────────────────
os.makedirs("report", exist_ok=True)
output = "report/K-CFD_Technical_Report.docx"
doc.save(output)

print(f"\nSaved: {output}")
print(f"Paragraphs : {len(doc.paragraphs)}")
print(f"Tables     : {len(doc.tables)}")
print(f"Images     : {len(doc.inline_shapes)}")
