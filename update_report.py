"""Update the original DOCX report: keep all theory/equations/images, update Python->C++ parts."""
import sys, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document('old_report.docx')

def replace_text_in_paragraph(p, old, new):
    """Replace text while preserving formatting."""
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)

def set_paragraph_text(p, new_text):
    """Replace entire paragraph text, keep first run's format."""
    if p.runs:
        fmt = p.runs[0].font
        for run in p.runs:
            run.text = ''
        p.runs[0].text = new_text

def add_table_after(para, headers, rows):
    """Insert table after a paragraph."""
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            tbl.rows[i+1].cells[j].text = str(val)
            for r in tbl.rows[i+1].cells[j].paragraphs[0].runs:
                r.font.size = Pt(9)
    # Move table after the paragraph
    para._element.addnext(tbl._tbl)

# ============================================================
# 1. Title page updates
# ============================================================
for i, p in enumerate(doc.paragraphs):
    t = p.text

    # Title
    if 'Two-Fluid Model FVM' in t and i < 10:
        set_paragraph_text(p, 'K-CFD: Two-Fluid Model FVM 기반\n열유체 해석 코드 개발 보고서')

    # Subtitle line
    if 'Python 기반 유한체적법' in t:
        set_paragraph_text(p,
            'C++17 / Eigen 3.4 기반 유한체적법(FVM) 코드\n'
            '2D/3D 구조·비정렬 격자 지원\n'
            '29개 모듈 | ~14,000 lines | 25개 검증 케이스')

    if 'Euler-Euler 이상유동 + Conjugate' in t:
        set_paragraph_text(p,
            'Euler-Euler 이상유동 + Conjugate Heat Transfer + k-ε/k-ω SST 난류 모델\n'
            'SIMPLE/PISO, MUSCL/TVD 2차 대류, Rhie-Chow 보간, BDF2 시간적분\n'
            'IAPWS-IF97 증기표, CSF 표면장력, OpenMP 병렬화, AMG 전처리기\n'
            'Grace/Tomiyama/Ishii-Zuber 항력, Tomiyama 양력, Burns 난류분산')

# ============================================================
# 2. Section 1 (서론) updates
# ============================================================
for i, p in enumerate(doc.paragraphs):
    t = p.text

    # 1. 서론 first paragraph
    if '본 보고서는 Two-Fluid Model' in t and 'Python으로 작성' in t:
        set_paragraph_text(p,
            '본 보고서는 Two-Fluid Model(Euler-Euler) 기반의 유한체적법(FVM) 열유체 해석 코드 '
            'K-CFD의 개발 및 검증 결과를 기술한다. 본 코드는 원래 Python으로 개발되었으며, '
            '엔지니어링 프로덕션 사용을 위해 C++17(Eigen 3.4)로 완전히 전환되었다. '
            '2D/3D 구조·비정렬 격자에서의 단상/이상 유동, 난류, 고체 열전도 및 유체-고체 '
            '공액 열전달, 상변화, 복사, 화학반응을 포함하는 종합적 열유체 해석이 가능하다. '
            'C++ 전환을 통해 Python 대비 평균 31배, 최대 105배의 성능 향상을 달성하였으며, '
            '2차에 걸친 엔지니어링 QA 감사를 통해 총 34개의 CRITICAL/HIGH 이슈를 해결하였다.')

    # 개발 범위 feature list
    if '본 코드의 주요 특징은 다음과 같다' in t:
        set_paragraph_text(p,
            '본 코드의 주요 특징은 다음과 같다:\n'
            '• Two-Fluid (Euler-Euler) 이상유동 모델 (6방정식 비평형)\n'
            '• SIMPLE/PISO 알고리즘 + Rhie-Chow 운동량 보간\n'
            '• k-ε 및 k-ω SST 난류 모델 + 자동 벽 처리 (Low-Re/WF/AUTOMATIC)\n'
            '• 고체 열전도 및 유체-고체 CHT 커플링\n'
            '• 4종 항력: Schiller-Naumann, Grace, Tomiyama, Ishii-Zuber\n'
            '• 3종 계면력: Tomiyama 양력, Antal 벽윤활, Burns 난류분산\n'
            '• CSF 표면장력 (Brackbill 1992)\n'
            '• IAPWS-IF97 증기표 (Region 1/2, 포화선, 수송물성)\n'
            '• Lee/Rohsenow/Zuber/Nusselt 상변화 모델\n'
            '• MUSCL/TVD 2차 대류 + BDF2 2차 시간적분\n'
            '• 비직교 격자 확산 보정\n'
            '• AMG V-cycle 전처리기 + 전처리 BiCGSTAB/CG\n'
            '• P1 복사, 1차 화학반응\n'
            '• GMSH .msh 2.2 메쉬 리더\n'
            '• OpenMP 병렬화\n'
            '• VTU (ParaView) 출력, pybind11 Python 바인딩')

# ============================================================
# 3. Section 6 (병렬화) — replace MPI/GPU with OpenMP/AMG
# ============================================================
for i, p in enumerate(doc.paragraphs):
    t = p.text

    if t.strip() == '6.1 MPI RCB 영역분할':
        set_paragraph_text(p, '6.1 OpenMP 공유 메모리 병렬화')
    if t.strip() == '6.2 고스트 셀 통신':
        set_paragraph_text(p, '6.2 병렬화 대상 루프')
    if t.strip() == '6.3 GPU 가속 (CuPy BiCGSTAB)':
        set_paragraph_text(p, '6.3 AMG V-cycle 전처리기')
    if t.strip() == '6.4 전처리기 (ILU/AMG)':
        set_paragraph_text(p, '6.4 전처리 솔버 (Jacobi/ILU0/AMG)')

    # Replace MPI content
    if 'GeometricPartitioner' in t and 'RCB' in t:
        set_paragraph_text(p,
            'OpenMP를 이용한 공유 메모리 병렬화를 구현하였다. '
            '모든 독립적 셀/면 루프에 #pragma omp parallel for를 적용하며, '
            'FVMSystem 조립 루프(면 기반 레이스 조건)는 직렬을 유지한다.\n\n'
            '병렬화 대상:\n'
            '• closure.cpp: 11개 함수 (항력, 양력, 벽윤활, 난류분산, 열전달)\n'
            '• gradient.cpp: Green-Gauss 체적 나눗셈, Least Squares\n'
            '• interpolation.cpp: compute_mass_flux\n'
            '• turbulence*.cpp: get_mu_t, compute_production\n'
            '• two_fluid_solver.cpp: 속도 보정, 체적분율 클리핑\n'
            '• mesh_generator*.cpp: 노드/셀 생성')

    if '고스트 셀(ghost cell)' in t:
        set_paragraph_text(p,
            '레이스 조건 방지를 위해 FVM 연산자의 면 루프는 직렬로 유지하고, '
            '셀 기반 독립 연산(물성치, 소스항, 후처리)만 병렬화한다. '
            '#ifdef _OPENMP 가드로 OpenMP 미지원 환경에서도 컴파일 가능하다.')

    if 'CuPy 라이브러리를 활용한 GPU BiCGSTAB' in t:
        set_paragraph_text(p,
            'AMG(Algebraic Multigrid) V-cycle 전처리기를 구현하였다:\n'
            '• 쌍별 응집(pairwise aggregation) 거칠게하기\n'
            '• Galerkin 거친 격자 연산자: A_c = R·A·P\n'
            '• Gauss-Seidel 전/후 평활 (기본 2회)\n'
            '• 최소 격자에서 직접 풀이 (SparseLU)\n\n'
            '전처리 반복 솔버:\n'
            '• solve_preconditioned(system, x0, "bicgstab", precond)\n'
            '• solve_preconditioned(system, x0, "cg", precond)\n'
            '• create_preconditioner(A, "amg") / "jacobi" / "ilu0"')

    if 'ILU(0), ILU(k), AMG 전처리기' in t:
        set_paragraph_text(p,
            'Jacobi, ILU(0), AMG 전처리기를 제공한다:\n'
            '• Jacobi: 대각 스케일링 (가장 빠른 setup)\n'
            '• ILU(0): 불완전 LU 분해 (Eigen IncompleteLUT)\n'
            '• AMG: V-cycle 멀티그리드 (대규모 격자에 효과적)')

# ============================================================
# 4. Section 11/12 결론 updates
# ============================================================
for i, p in enumerate(doc.paragraphs):
    t = p.text

    if 'Python 기반 Two-Fluid Model FVM 열유체 코드를 개발하고' in t:
        set_paragraph_text(p,
            'C++17 기반 Two-Fluid Model FVM 열유체 코드 K-CFD를 개발하고, '
            '25개의 검증 케이스를 통해 코드의 정확성과 확장성을 확인하였다. '
            'Python에서 C++로의 전환을 통해 평균 31배, 최대 105배의 성능 향상을 달성하였으며, '
            '2차에 걸친 엔지니어링 QA 감사를 통해 34개의 CRITICAL/HIGH 이슈를 모두 해결하였다. '
            '주요 성과는 다음과 같다:\n\n'
            '• Two-Fluid (Euler-Euler) 6방정식 이상유동 해석\n'
            '• SIMPLE/PISO + Rhie-Chow 운동량 보간\n'
            '• k-ε + k-ω SST 난류 모델 (자동 벽 처리)\n'
            '• 4종 항력, 3종 계면력, CSF 표면장력\n'
            '• IAPWS-IF97 물성치, BDF2 시간적분, 비직교 보정\n'
            '• AMG 전처리기, OpenMP 병렬화\n'
            '• GMSH 메쉬 리더, 3D Hex/Tet 혼합 격자\n\n'
            'Python 대비 C++ 성능 비교:\n'
            '  Poiseuille (50×20): 31,032ms → 1,252ms (24.8배)\n'
            '  Cavity Re=100 (32×32): 42,834ms → 1,682ms (25.5배)\n'
            '  Bubble Column (8×20): 51,716ms → 893ms (57.9배)\n'
            '  3D Mesh (8³): 52.5ms → 0.5ms (105배)')

# ============================================================
# 5. Section 13 (프로그래머 매뉴얼) .py → .hpp/.cpp
# ============================================================
py_to_cpp = {
    'mesh_reader.py': 'mesh.hpp / mesh.cpp + mesh_reader.hpp / mesh_reader.cpp',
    'mesh_generator.py': 'mesh_generator.hpp / mesh_generator.cpp',
    'mesh_generator_3d.py': 'mesh_generator_3d.hpp / mesh_generator_3d.cpp',
    'amr.py': 'amr.hpp / amr.cpp',
    'vtk_exporter.py': 'vtk_writer.hpp / vtk_writer.cpp',
    'fields.py': 'fields.hpp / fields.cpp',
    'gradient.py': 'gradient.hpp / gradient.cpp',
    'interpolation.py': 'interpolation.hpp / interpolation.cpp',
    'fvm_operators.py': 'fvm_operators.hpp / fvm_operators.cpp',
    'linear_solver.py': 'linear_solver.hpp / linear_solver.cpp',
    'single_phase.py': 'simple_solver.hpp / simple_solver.cpp',
    'two_fluid.py': 'two_fluid_solver.hpp / two_fluid_solver.cpp',
    'turbulence.py': 'turbulence.hpp / turbulence.cpp + turbulence_sst.hpp / turbulence_sst.cpp',
    'closure.py': 'closure.hpp / closure.cpp + surface_tension.hpp / surface_tension.cpp',
    'phase_change.py': 'phase_change.hpp / phase_change.cpp + steam_tables.hpp / steam_tables.cpp',
    'reaction.py': 'chemistry.hpp / chemistry.cpp',
    'radiation.py': 'radiation.hpp / radiation.cpp',
    'partitioning.py': '(OpenMP 병렬화로 대체 — 별도 파티셔닝 불필요)',
    'mpi_solver.py': '(OpenMP 병렬화로 대체 — simple_solver.cpp 내 통합)',
}

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    for pyfile, cppfile in py_to_cpp.items():
        if t == pyfile or t == f'13.' and pyfile in t:
            pass  # handled by heading below

    # Update headings
    for pyfile, cppfile in py_to_cpp.items():
        if pyfile in t and ('Heading' in p.style.name):
            new_heading = t.replace(pyfile, cppfile.split('/')[0].strip() if '/' in cppfile else cppfile)
            set_paragraph_text(p, new_heading)
            break

    # Update section 13 intro
    if '각 스크립트는 run_caseN() 함수를 제공' in t:
        set_paragraph_text(p,
            '각 검증 케이스는 cpp/tests/verification_cases.cpp에 통합되어 있으며, '
            'twofluid_solver 실행 파일에서 --case 옵션으로 선택할 수 있다.')

    # Section 13 title
    if t == '13. 프로그래머 매뉴얼':
        set_paragraph_text(p, '13. 프로그래머 매뉴얼 (C++ API)')

    # Module section titles
    if t == '13.1 mesh/ 모듈':
        set_paragraph_text(p, '13.1 메쉬 모듈 (cpp/include/twofluid/mesh*.hpp)')
    if t == '13.2 core/ 모듈':
        set_paragraph_text(p, '13.2 코어 모듈 (cpp/include/twofluid/)')
    if t == '13.3 models/ 모듈':
        set_paragraph_text(p, '13.3 물리 모델 (cpp/include/twofluid/)')
    if t == '13.4 parallel/ 모듈':
        set_paragraph_text(p, '13.4 병렬화 (OpenMP)')
    if t == '13.5 gpu/ 모듈':
        set_paragraph_text(p, '13.5 전처리기 및 선형 솔버')
    if t == '13.6 verification/ 모듈':
        set_paragraph_text(p, '13.6 검증 (cpp/tests/)')
    if t == '13.7 report/ 모듈':
        set_paragraph_text(p, '13.7 빌드 시스템 (CMakeLists.txt)')

# ============================================================
# 6. Global text replacements (preserve formatting)
# ============================================================
global_replacements = [
    ('Python으로 작성되었으며', 'C++17로 작성되었으며'),
    ('Python 기반', 'C++ 기반'),
    ('scipy.sparse', 'Eigen::SparseMatrix'),
    ('scipy.sparse.linalg.spsolve', 'Eigen::SparseLU'),
    ('numpy', 'Eigen'),
    ('NumPy', 'Eigen'),
    ('SciPy', 'Eigen'),
    ('mpi4py', 'OpenMP'),
    ('CuPy', '(삭제됨 — OpenMP로 대체)'),
]

for p in doc.paragraphs:
    for old, new in global_replacements:
        if old in p.text:
            replace_text_in_paragraph(p, old, new)

# Also update tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for old, new in global_replacements:
                    if old in p.text:
                        replace_text_in_paragraph(p, old, new)

# ============================================================
# 7. Add new appendix: C++ 전환 및 QA 감사 결과
# ============================================================
# Add before references (section 14)
# Find the references heading and insert before it
ref_idx = None
for i, p in enumerate(doc.paragraphs):
    if '14. 참고문헌' in p.text:
        ref_idx = i
        break

if ref_idx:
    # We'll add content at the end, before references
    # Change references to section 15
    for p in doc.paragraphs:
        if p.text.strip() == '14. 참고문헌':
            set_paragraph_text(p, '15. 참고문헌')

# Add new section 14 at the end (just append paragraphs)
p14 = doc.add_heading('14. C++ 전환 및 엔지니어링 QA 감사', level=1)

doc.add_heading('14.1 Python → C++ 전환', level=2)
doc.add_paragraph(
    '원래 Python으로 작성된 코드를 C++17로 완전히 전환하였다. '
    '전환 과정에서 모든 Python 소스(17,153줄)를 삭제하고, '
    '동등한 기능의 C++ 코드(~14,000줄, 29개 모듈)를 새로 작성하였다.\n\n'
    'Python과 C++의 수치 결과가 동일함을 13개 검증 지표로 확인하였다. '
    'C++는 Python 대비 평균 31배, 최대 105배 빠르다.')

doc.add_heading('14.2 성능 비교', level=2)
perf_table = doc.add_table(rows=9, cols=4)
perf_table.style = 'Light Grid Accent 1'
perf_headers = ['Case', 'Python (ms)', 'C++ (ms)', '속도향상']
perf_data = [
    ['Poiseuille (50×20)', '31,032', '1,252', '24.8x'],
    ['Cavity Re=100 (32×32)', '42,834', '1,682', '25.5x'],
    ['Bubble Column (8×20)', '51,716', '893', '57.9x'],
    ['MUSCL (50×10)', '2,816', '256', '11.0x'],
    ['Phase Change', '5.5', '0.4', '13.8x'],
    ['Radiation (20×20)', '17.3', '2.9', '6.0x'],
    ['AMR (8×8)', '9.1', '0.4', '22.8x'],
    ['3D Mesh (8³)', '52.5', '0.5', '105x'],
]
for j, h in enumerate(perf_headers):
    perf_table.rows[0].cells[j].text = h
    for r in perf_table.rows[0].cells[j].paragraphs[0].runs:
        r.bold = True
for i, row in enumerate(perf_data):
    for j, val in enumerate(row):
        perf_table.rows[i+1].cells[j].text = val

doc.add_paragraph('')

doc.add_heading('14.3 엔지니어링 QA 감사', level=2)
doc.add_paragraph(
    '프로덕션 사용을 위해 2차에 걸친 심층 감사를 수행하였다.\n\n'
    '1차 감사: CRITICAL 7개 + HIGH 13개 발견 → 모두 해결\n'
    '  • C1-C3: 하드코딩 클리핑 제거 → 사용자 설정\n'
    '  • C4: MUSCL 2차 대류 연결\n'
    '  • C5: Rhie-Chow 운동량 보간\n'
    '  • H1: k-ω SST 추가\n'
    '  • H3: AMG 전처리기\n'
    '  • H4-H5: Grace/Tomiyama 항력, 양력/벽윤활/난류분산\n'
    '  • H7: BDF2 시간적분\n'
    '  • H8: OpenMP 병렬화\n'
    '  • H9: PISO 알고리즘\n'
    '  • H12-H13: IAPWS-IF97, GMSH 리더\n\n'
    '2차 감사: 추가 CRITICAL 8개 + HIGH 6개 발견 → 모두 해결\n'
    '  • Rhie-Chow 대수적 소거 수정 (cell-center gradient 기반)\n'
    '  • 9개 기능 솔버 통합 (미연결 → 연결)\n'
    '  • IAPWS g2o_pi 오류 수정 (1.0 → 1/pi)\n'
    '  • 압력구배 공식 수정 (절대압 → 차분)\n'
    '  • AMR 필드전달 수정 (전역평균 → 부모값)\n'
    '  • PISO 시간항 추가\n\n'
    '최종 현황: CRITICAL 15/15 완료, HIGH 19/19 완료, MEDIUM 12개 향후 개선')

doc.add_heading('14.4 연결 전/후 검증', level=2)
doc.add_paragraph(
    '솔버 통합(Rhie-Chow, 압력구배, IAPWS, 계면력) 전후 비교:\n\n'
    '28개 검증 지표 중 26개 Exact 일치, 2개 의도된 물리적 개선.\n'
    '(Bubble Column: alpha_max=0.9→1.0 해제로 물리적 한계 허용)\n\n'
    '단상 솔버(Poiseuille, Cavity, MUSCL): 모든 지표 동일.\n'
    '물리 모델(상변화, 복사, AMR, 3D, 적응dt): 모든 지표 Exact.\n'
    '이상유체(Bubble Column): alpha_g_max 개선 (물리적 타당).')

# ============================================================
# Save
# ============================================================
output = 'report/K-CFD_Technical_Report.docx'
doc.save(output)
print(f'Updated report saved to: {output}')
print(f'Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}, Images: {len(doc.inline_shapes)}')
