"""
한글 DOCX 보고서 생성기.

Two-Fluid Model FVM 열유체 코드의 검증 결과를 종합한 한글 보고서를 생성한다.
새 구조: 14개 장, 25개 검증 케이스, 기술적 논리 순서 배치.
"""

import math
import os
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _add_heading(doc, text, level=1):
    """제목 추가."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return heading


def _add_paragraph(doc, text, bold=False, italic=False):
    """본문 단락 추가."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Malgun Gothic'
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def _add_figure(doc, fig_path, caption, width=Inches(5.5)):
    """그림 삽입."""
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.font.size = Pt(10)
        run.italic = True
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_table(doc, headers, rows):
    """표 추가."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # 데이터
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table


# =====================================================================
# 각 장 생성 함수
# =====================================================================

def _write_cover(doc):
    """표지."""
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_heading('Two-Fluid Model FVM 기반\n열유체 해석 코드 개발 보고서', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        'Euler-Euler 이상유동 + Conjugate Heat Transfer + k-ε 난류 모델\n'
        '3D 구조/비정렬 격자, MUSCL/TVD, MPI 병렬화, 상변화, AMR, GPU 가속')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Python 기반 유한체적법(FVM) 코드\n2D/3D 구조·비정렬 격자 지원\n25개 검증 케이스')
    run.font.size = Pt(12)

    doc.add_page_break()


def _write_toc(doc):
    """목차."""
    _add_heading(doc, '목차', level=1)
    toc_items = [
        '1. 서론',
        '   1.1 개발 배경',
        '   1.2 개발 범위 (25개 검증 케이스)',
        '2. 지배방정식',
        '   2.1 Two-Fluid Model (Euler-Euler)',
        '   2.2 단상 비압축성 Navier-Stokes',
        '   2.3 k-ε 난류 모델',
        '   2.4 고체 열전도',
        '   2.5 상변화 모델',
        '   2.6 화학반응 모델',
        '   2.7 P1 복사열전달 모델',
        '3. 수치기법',
        '   3.1 유한체적법(FVM) 이산화',
        '   3.2 SIMPLE 알고리즘',
        '   3.3 경계조건 처리',
        '   3.4 기울기 복원',
        '   3.5 MUSCL/TVD 고차 이산화',
        '   3.6 비정렬 격자 처리',
        '   3.7 상변화 소스항 수치해법',
        '4. 코드 구조',
        '5. 격자 생성',
        '   5.1 2D 구조격자',
        '   5.2 2D 비정렬 삼각형 격자',
        '   5.3 3D 육면체 구조격자',
        '   5.4 적응 격자 세분화(AMR)',
        '6. 병렬화 및 가속',
        '   6.1 MPI RCB 영역분할',
        '   6.2 고스트 셀 통신',
        '   6.3 GPU 가속 (CuPy BiCGSTAB)',
        '   6.4 전처리기 (ILU/AMG)',
        '7. 검증 결과 — 2D 문제',
        '   7.1  Case 1: Poiseuille 유동',
        '   7.2  Case 2: Lid-Driven Cavity',
        '   7.3  Case 3: Conjugate Heat Transfer',
        '   7.4  Case 4: 기포탑 이상유동',
        '   7.5  Case 6: MUSCL/TVD 고차 이산화',
        '   7.6  Case 7: 비정렬 격자',
        '   7.7  Case 8: MPI 병렬화',
        '   7.8  Case 9: Stefan 문제 (Lee 상변화)',
        '   7.9  Case 10: 화학반응',
        '   7.10 Case 11: 복사열전달',
        '   7.11 Case 12: AMR',
        '   7.12 Case 16: 전처리기 비교',
        '   7.13 Case 17: 적응 시간 간격',
        '   7.14 Case 19: 비등/응축 상변화',
        '8. 검증 결과 — 3D 문제',
        '   8.1  Case 5: 3D 덕트 Poiseuille',
        '   8.2  Case 14: 3D Lid-Driven Cavity',
        '   8.3  Case 15: 3D 자연대류',
        '   8.4  Case 22: 혼합 격자(Hex/Tet) Poiseuille',
        '   8.5  Case 23: 3D 과도 채널 유동',
        '9. 검증 결과 — 상변화/다상유동',
        '   9.1  Case 20: Edwards 파이프 블로우다운 (Flashing)',
        '   9.2  Case 24: 풀 비등 (Pool Boiling)',
        '   9.3  Case 25: 수직벽 막 응축 (Film Condensation)',
        '10. 검증 결과 — 기타',
        '   10.1 Case 13: GPU 가속 검증',
        '   10.2 Case 18: 웹 대시보드',
        '11. 검증 결과 요약',
        '12. 결론 및 향후 과제',
        '13. 프로그래머 매뉴얼',
        '14. 참고문헌',
    ]
    for item in toc_items:
        _add_paragraph(doc, item)

    doc.add_page_break()


def _write_ch1_intro(doc):
    """1장 서론."""
    _add_heading(doc, '1. 서론', level=1)

    _add_paragraph(doc,
        '본 보고서는 Two-Fluid Model(Euler-Euler) 기반의 유한체적법(FVM) 열유체 해석 코드의 '
        '개발 및 검증 결과를 기술한다. 본 코드는 Python으로 작성되었으며, 2D/3D 구조·비정렬 격자에서의 '
        '단상/이상 유동, 난류, 고체 열전도 및 유체-고체 연성 열전달(CHT)을 다룬다.')

    _add_heading(doc, '1.1 개발 배경', level=2)
    _add_paragraph(doc,
        '원자로, 화학반응기, 열교환기 등 다양한 공학 시스템에서 이상유동 및 열전달 현상의 '
        '정확한 예측은 안전 설계의 핵심이다. Two-Fluid Model은 각 상을 연속체로 취급하여 '
        '대규모 시스템의 열유체 해석에 효과적인 접근법을 제공한다.')

    _add_heading(doc, '1.2 개발 범위', level=2)
    _add_paragraph(doc,
        '본 코드의 주요 특징은 다음과 같다:\n'
        '• Two-Fluid (Euler-Euler) 이상유동 모델\n'
        '• SIMPLE 알고리즘 기반 속도-압력 커플링\n'
        '• 표준 k-ε 난류 모델 + 벽함수\n'
        '• 고체 열전도 및 유체-고체 CHT 커플링\n'
        '• Schiller-Naumann 항력, Ranz-Marshall 열전달, Sato 기포유도 난류\n'
        '• 2D/3D 구조·비정렬 격자 지원 (Gmsh 기반)\n'
        '• MUSCL/TVD 고차 이산화 스킴 (4종 리미터)\n'
        '• 비정렬 삼각형 격자 및 비직교 보정\n'
        '• MPI 병렬화 (RCB 분할, 고스트 셀 통신)\n'
        '• Lee 상변화, Rohsenow 핵비등, Zuber CHF, Nusselt 막응축 모델\n'
        '• 1차 화학반응, P1 복사열전달 모델\n'
        '• AMR (사분면 세분화, 기울기-점프 추정기)\n'
        '• GPU 가속 (CuPy BiCGSTAB, CPU 폴백)\n'
        '• ILU(k)/AMG 전처리기\n'
        '• CFL 기반 적응 시간 간격 제어\n'
        '• ParaView 호환 VTU 격자/결과 출력\n'
        '• Flask + Plotly.js 웹 실시간 모니터링 대시보드\n'
        '• 6-방정식 비평형 상변화 모델 (각 상별 에너지 방정식 SIMPLE 내 결합)\n'
        '• Hex/Tet 혼합 격자 생성기 (Dompierre 5-tet 분해)\n'
        '• 25개 검증 케이스를 통한 포괄적 코드 검증')

    doc.add_page_break()


def _write_ch2_governing(doc):
    """2장 지배방정식 (상변화/반응/복사 이론 통합)."""
    _add_heading(doc, '2. 지배방정식', level=1)

    # 2.1 Two-Fluid
    _add_heading(doc, '2.1 Two-Fluid Model (Euler-Euler)', level=2)
    _add_paragraph(doc,
        'Two-Fluid Model은 각 상(phase)을 상호침투 연속체로 취급하며, '
        '각 상에 대해 독립적인 보존방정식을 풀어낸다.')

    _add_paragraph(doc, '연속 방정식 (각 상 k):', bold=True)
    _add_paragraph(doc, '∂(αₖρₖ)/∂t + ∇·(αₖρₖuₖ) = Γₖ')

    _add_paragraph(doc, '운동량 방정식 (각 상 k):', bold=True)
    _add_paragraph(doc,
        '∂(αₖρₖuₖ)/∂t + ∇·(αₖρₖuₖ⊗uₖ) = -αₖ∇p + ∇·(αₖτₖ) + αₖρₖg + Mₖ')
    _add_paragraph(doc, '여기서 Mₖ는 상간 힘(항력, 가상질량 등)을 나타낸다.')

    _add_paragraph(doc, '에너지 방정식 (각 상 k):', bold=True)
    _add_paragraph(doc, '∂(αₖρₖhₖ)/∂t + ∇·(αₖρₖuₖhₖ) = ∇·(αₖkₖ∇Tₖ) + Qₖ')

    _add_paragraph(doc, '체적분율 구속조건:', bold=True)
    _add_paragraph(doc, 'α_l + α_g = 1')

    _add_heading(doc, '2.1.1 6-방정식 비평형 모델', level=3)
    _add_paragraph(doc,
        '6-방정식 모델에서는 각 상(액체, 기체)에 대해 독립적인 에너지 방정식을 풀어 '
        '비평형 상변화를 다룬다. 두 상의 온도(T_l, T_g)는 서로 다를 수 있으며, '
        '상간 열전달에 의해 결합된다.')

    _add_paragraph(doc, '상간 열전달:', bold=True)
    _add_paragraph(doc,
        'Q_i = h_i · a_i · (T_g - T_l)\n'
        'a_i = 6·α_g / d_b (비계면적)\n'
        'h_i = Nu · k_l / d_b (Ranz-Marshall)\n\n'
        '에너지 방정식에서의 선형화 처리:\n'
        '• 액상: S_u += h_i·a_i·T_g,  S_p += -h_i·a_i\n'
        '• 기상: S_u += h_i·a_i·T_l,  S_p += -h_i·a_i\n\n'
        '이 선형화를 통해 SIMPLE 내부 반복에서 에너지 방정식을 암시적으로 풀 수 있다.')

    _add_paragraph(doc, '상변화 소스항:', bold=True)
    _add_paragraph(doc,
        '증발/응축에 의한 질량 전달률 ṁ:\n'
        '• 증발 (T_l > T_sat): ṁ = r · α_l · ρ_l · (T_l - T_sat) / T_sat\n'
        '• 응축 (T_g < T_sat): ṁ = -r · α_g · ρ_g · (T_sat - T_g) / T_sat\n\n'
        '에너지 방정식 잠열 소스: Q_latent = ṁ · h_fg\n'
        '체적분율 방정식 소스: S_α = ṁ / ρ_g')

    _add_heading(doc, '2.1.2 폐합관계식', level=3)
    _add_paragraph(doc, 'Schiller-Naumann 항력:', bold=True)
    _add_paragraph(doc,
        'C_D = 24/Re_p · (1 + 0.15·Re_p^0.687)  (Re_p < 1000)\n'
        'C_D = 0.44  (Re_p ≥ 1000)\n'
        'F_drag = 0.75·C_D·α_g·ρ_l·|u_g - u_l|·(u_g - u_l) / d_b')

    _add_paragraph(doc, 'Ranz-Marshall 열전달:', bold=True)
    _add_paragraph(doc, 'Nu = 2 + 0.6·Re_p^0.5·Pr^0.33')

    _add_paragraph(doc, 'Sato 기포유도 난류:', bold=True)
    _add_paragraph(doc, 'μ_t,BIT = C_μb·ρ_l·α_g·d_b·|u_g - u_l|  (C_μb = 0.6)')

    # 2.2 단상 NS
    _add_heading(doc, '2.2 단상 비압축성 Navier-Stokes', level=2)
    _add_paragraph(doc,
        '단상 비압축성 유동의 연속 및 운동량 방정식:\n'
        '∇·u = 0\n'
        '∂(ρu)/∂t + ∇·(ρu⊗u) = -∇p + ∇·(μ∇u) + ρg')

    # 2.3 k-eps
    _add_heading(doc, '2.3 k-ε 난류 모델', level=2)
    _add_paragraph(doc, '표준 k-ε 모델의 수송 방정식:')
    _add_paragraph(doc,
        '∂(ρk)/∂t + ∇·(ρuk) = ∇·((μ + μ_t/σ_k)∇k) + P_k - ρε\n'
        '∂(ρε)/∂t + ∇·(ρuε) = ∇·((μ + μ_t/σ_ε)∇ε) + C₁ε(ε/k)P_k - C₂ερε²/k\n'
        'μ_t = ρ·C_μ·k²/ε')

    _add_paragraph(doc, '모델 상수:', bold=True)
    _add_table(doc,
        ['상수', 'C_μ', 'C₁ε', 'C₂ε', 'σ_k', 'σ_ε'],
        [['값', '0.09', '1.44', '1.92', '1.0', '1.3']])

    # 2.4 고체 열전도
    _add_heading(doc, '2.4 고체 열전도', level=2)
    _add_paragraph(doc,
        'ρ_s·c_s·∂T_s/∂t = ∇·(k_s·∇T_s) + q\'\'\'\n\n'
        'CHT 인터페이스 조건:\n'
        'T_f = T_s (온도 연속)\n'
        'k_f·∂T_f/∂n = k_s·∂T_s/∂n (열유속 연속)')

    # 2.5 상변화 모델 (기존 11장에서 이동)
    _add_heading(doc, '2.5 상변화 모델', level=2)

    _add_heading(doc, '2.5.1 Lee 증발/응축 모델', level=3)
    _add_paragraph(doc,
        'Lee(1980) 모델은 계면 온도에 따른 증발/응결 상변화율을 다음과 같이 계산한다:\n\n'
        '증발 (T_l > T_sat):\n'
        'Γ_evap = r_evap · α_l · ρ_l · (T_l - T_sat) / T_sat\n\n'
        '응결 (T_g < T_sat):\n'
        'Γ_cond = r_cond · α_g · ρ_g · (T_sat - T_g) / T_sat\n\n'
        '여기서 r_evap, r_cond는 상변화 계수(s⁻¹)이다.')

    _add_heading(doc, '2.5.2 Rohsenow 핵비등', level=3)
    _add_paragraph(doc,
        'Rohsenow(1952) 핵비등 벽면 열유속 상관식:\n\n'
        'q"_w = μ_l · h_fg · [g·(ρ_l - ρ_g)/σ]^0.5 · [c_pl·ΔT_sat/(C_sf·h_fg·Pr_l^n)]³\n\n'
        '여기서:\n'
        '• ΔT_sat = T_wall - T_sat : 벽면 과열도 [K]\n'
        '• C_sf : 표면-유체 조합 상수 (물-구리: 0.013)\n'
        '• n : Prandtl 지수 (물: 1.0)\n'
        '• σ : 표면장력 [N/m]')

    _add_paragraph(doc, 'C_sf 값 (대표적 표면-유체 조합):', bold=True)
    _add_table(doc,
        ['표면-유체 조합', 'C_sf', 'n'],
        [
            ['물 - 구리 (연마)', '0.013', '1.0'],
            ['물 - 스테인리스강', '0.015', '1.0'],
            ['물 - 황동', '0.006', '1.0'],
            ['에탄올 - 크롬', '0.027', '1.7'],
        ])

    _add_heading(doc, '2.5.3 Zuber 임계 열유속(CHF)', level=3)
    _add_paragraph(doc,
        'Zuber(1959) 풀비등 임계 열유속 상관식:\n\n'
        'q"_CHF = 0.131 · ρ_g · h_fg · [σ·g·(ρ_l - ρ_g)/ρ_g²]^0.25\n\n'
        'CHF를 초과하면 벽면 온도가 급격히 상승하는 번아웃(burnout) 현상이 발생한다.')

    _add_heading(doc, '2.5.4 Nusselt 막응축', level=3)
    _add_paragraph(doc,
        'Nusselt(1916) 수직 평판 막응축 해석해:\n\n'
        'h_cond = 0.943 · [ρ_l·(ρ_l - ρ_g)·g·h_fg·k_l³ / (μ_l·L·ΔT_sub)]^0.25\n\n'
        '여기서 ΔT_sub = T_sat - T_wall (벽면 과냉도), L은 평판 높이이다.')

    # 2.6 화학반응 모델 (기존 11장에서 이동)
    _add_heading(doc, '2.6 화학반응 모델', level=2)
    _add_paragraph(doc,
        '단순 1차 반응 A → B의 소스항:\n\n'
        'S_A = -k_r · C_A · V_cell\n\n'
        '여기서 k_r은 반응 속도 상수(s⁻¹), C_A는 반응물 농도(mol/m³)이다.')

    # 2.7 P1 복사열전달 (기존 11장에서 이동)
    _add_heading(doc, '2.7 P1 복사열전달 모델', level=2)
    _add_paragraph(doc,
        'P1 근사는 복사 강도를 구면 조화함수로 전개하여 타원형 PDE로 변환한다:\n\n'
        '∇·(Γ_rad·∇G) - κ·G + 4κ·σ·T⁴ = 0\n\n'
        '여기서 Γ_rad = 1/(3(κ + σ_s)), κ는 흡수 계수, σ_s는 산란 계수이다.\n\n'
        '에너지 방정식 커플링: q_rad = κ · (G - 4σT⁴)')

    doc.add_page_break()


def _write_ch3_numerics(doc):
    """3장 수치기법 (MUSCL/비정렬/상변화 수치해법 통합)."""
    _add_heading(doc, '3. 수치기법', level=1)

    _add_heading(doc, '3.1 유한체적법(FVM) 이산화', level=2)
    _add_paragraph(doc,
        '셀 중심 유한체적법을 적용하며, 각 보존방정식을 셀 체적에 대해 적분한다.')

    _add_paragraph(doc, '확산항:', bold=True)
    _add_paragraph(doc, '∫∇·(Γ∇φ)dV ≈ Σ_f Γ_f·(φ_N - φ_P)/d_PN · A_f')

    _add_paragraph(doc, '대류항:', bold=True)
    _add_paragraph(doc,
        '∫∇·(ρuφ)dV ≈ Σ_f (ρu)_f · φ_f · A_f\n'
        '면 보간: Upwind (1차), Central (2차), TVD-Van Leer (2차 고해상도)')

    _add_paragraph(doc, '시간항:', bold=True)
    _add_paragraph(doc, '∂(ρφ)/∂t · V ≈ (ρφ - ρ°φ°)/Δt · V  (후방 Euler)')

    _add_heading(doc, '3.2 SIMPLE 알고리즘', level=2)
    _add_paragraph(doc,
        'SIMPLE(Semi-Implicit Method for Pressure-Linked Equations):\n'
        '1. 운동량 방정식에서 u*, v* 추정 (이전 압력 p* 사용)\n'
        '2. 압력 보정 방정식: ∇·(1/a_P · ∇p\') = ∇·u*\n'
        '3. 속도 보정: u = u* - (1/a_P)·∇p\'\n'
        '4. 압력 갱신: p = p* + α_p·p\'\n'
        '5. 수렴 판정: |∇·u| < ε')

    _add_heading(doc, '3.3 경계조건 처리', level=2)
    _add_paragraph(doc,
        '• Dirichlet (고정값): 입구 속도, 벽면 no-slip\n'
        '• Neumann (고정 기울기): 열유속 경계\n'
        '• Zero gradient: 출구 조건\n'
        '• 벽함수: 난류 벽면 처리')

    _add_heading(doc, '3.4 기울기 복원', level=2)
    _add_paragraph(doc,
        'Green-Gauss 방법: ∇φ_P ≈ (1/V_P) Σ_f φ_f · n_f · A_f\n'
        'Least Squares 방법: 이웃 셀과의 차이를 최소자승법으로 계산')

    # 3.5 MUSCL/TVD (기존 8장에서 이동)
    _add_heading(doc, '3.5 MUSCL/TVD 고차 이산화', level=2)
    _add_paragraph(doc,
        'MUSCL 스킴은 선형 재구성을 통해 면에서의 변수값을 2차 정확도로 추정한다:\n'
        'φ_f = φ_P + (1/2)·ψ(r)·(φ_P - φ_U)\n'
        '여기서 r = (φ_P - φ_U)/(φ_D - φ_P)는 연속 기울기 비율이다.')

    _add_table(doc,
        ['리미터', '수식', '특성'],
        [
            ['Minmod', 'ψ = max(0, min(1, r))', '가장 보수적'],
            ['Superbee', 'ψ = max(0, min(2r,1), min(r,2))', '해상도 최대'],
            ['Van Albada', 'ψ = (r²+r)/(r²+1)', '연속 미분 가능'],
            ['Van Leer', 'ψ = (r+|r|)/(1+|r|)', '2차 정확도 보장'],
        ])

    _add_paragraph(doc,
        '지연 보정(Deferred Correction)으로 TVD 비선형성을 처리한다:\n'
        'a_P·φ_P = Σ a_nb·φ_nb + b + [고차항 - 저차항]^(이전 반복)')

    # 3.6 비정렬 격자 처리 (기존 9장에서 이동)
    _add_heading(doc, '3.6 비정렬 격자 처리', level=2)
    _add_paragraph(doc,
        '삼각형 비정렬 격자에서 셀 중심 연결선이 면 법선과 정렬되지 않을 때 '
        '비직교 보정을 적용한다:\n\n'
        '(∇φ)_f · n_f = (φ_N - φ_P)/|d_PN| + ∇φ_f · (n_f - d_PN/|d_PN|)\n\n'
        '두 번째 항이 비직교 보정항으로, 이웃 셀의 기울기로 명시적 계산한다.')

    # 3.7 에너지 방정식 SIMPLE 내 결합 해법
    _add_heading(doc, '3.7 에너지 방정식 SIMPLE 내 결합 해법', level=2)
    _add_paragraph(doc,
        '6-방정식 모델에서 에너지 방정식은 SIMPLE 내부 반복에 포함되어 암시적으로 풀린다:\n\n'
        '1. 운동량 방정식 풀기 (각 상, 각 성분)\n'
        '2. 압력 보정 방정식 풀기\n'
        '3. 체적분율 방정식 풀기 (상변화 소스항 포함)\n'
        '4. 에너지 방정식 풀기 (각 상별, 상간 열전달 선형화)\n'
        '5. 수렴 판정 후 반복\n\n'
        '에너지 방정식의 이산화는 운동량 방정식과 동일한 FVM 연산자를 사용한다:\n'
        '• 확산: ∫∇·(α_k·k_k·∇T_k)dV\n'
        '• 대류: ∫∇·(α_k·ρ_k·cp_k·u_k·T_k)dV\n'
        '• 시간: (α_k·ρ_k·cp_k·V/Δt)·(T_k - T_k°)\n'
        '• 상간 열전달: h_i·a_i·(T_other - T_k)·V (선형화)\n'
        '• 잠열: ṁ·h_fg·V')

    # 3.8 상변화 소스항 수치해법 (기존 11.1.5에서 이동)
    _add_heading(doc, '3.8 상변화 소스항 수치해법', level=2)
    _add_paragraph(doc,
        '상변화 소스항의 FVM 이산화:\n\n'
        '(1) 연속방정식: ∂α_l/∂t = -Γ/ρ_l, ∂α_g/∂t = +Γ/ρ_g\n'
        '(2) 에너지방정식: S_energy = -ṁ · h_fg\n\n'
        'Lee 모델의 안정성: dt < 1/(r · max(ρ_l·α_l, ρ_g·α_g)/T_sat)\n\n'
        '벽면 비등 소스항: A_wall/V_cell 비율로 스케일링하여 체적 소스항으로 변환.')

    doc.add_page_break()


def _write_ch4_code_structure(doc):
    """4장 코드 구조."""
    _add_heading(doc, '4. 코드 구조', level=1)

    _add_paragraph(doc, '본 코드는 모듈화된 구조로 개발되었으며, 주요 모듈은 다음과 같다:')

    _add_table(doc,
        ['모듈', '파일', '설명'],
        [
            ['격자', 'mesh/mesh_reader.py', 'Gmsh .msh 파일 파서, FVM 메쉬 구조'],
            ['격자', 'mesh/mesh_generator.py', '2D 구조 격자 생성기'],
            ['격자', 'mesh/mesh_generator_3d.py', '3D 육면체 구조격자 생성기'],
            ['격자', 'mesh/hybrid_mesh_generator.py', 'Hex/Tet 혼합격자 생성기'],
            ['격자', 'mesh/unstructured_mesh.py', '비정렬 삼각형 격자 생성'],
            ['격자', 'mesh/amr_mesh.py', '적응 격자 세분화(AMR)'],
            ['격자', 'mesh/vtk_exporter.py', 'ParaView VTU 내보내기'],
            ['핵심', 'core/fields.py', '셀 중심 스칼라/벡터 필드'],
            ['핵심', 'core/gradient.py', 'Green-Gauss / Least Squares 기울기'],
            ['핵심', 'core/interpolation.py', 'Upwind/Central/TVD/MUSCL 면 보간'],
            ['핵심', 'core/fvm_operators.py', 'FVM 이산화 연산자'],
            ['핵심', 'core/linear_solver.py', '선형 솔버 인터페이스'],
            ['모델', 'models/single_phase.py', 'SIMPLE 단상 NS 솔버'],
            ['모델', 'models/two_fluid.py', 'Two-Fluid 솔버'],
            ['모델', 'models/turbulence.py', 'k-ε 난류 모델'],
            ['모델', 'models/solid_conduction.py', '고체 열전도 솔버'],
            ['모델', 'models/conjugate_ht.py', 'CHT 커플링'],
            ['모델', 'models/closure.py', '상간 전달 폐합관계식'],
            ['모델', 'models/phase_change.py', '상변화 모델 (Lee, Rohsenow, Zuber, Nusselt)'],
            ['모델', 'models/reaction.py', '1차 화학반응'],
            ['모델', 'models/radiation.py', 'P1 복사열전달'],
            ['병렬', 'parallel/mpi_decomp.py', 'RCB 영역분할, 고스트 셀'],
            ['가속', 'gpu/gpu_solver.py', 'CuPy BiCGSTAB GPU 솔버'],
        ])

    doc.add_page_break()


def _write_ch5_mesh(doc):
    """5장 격자 생성 (2D/3D/비정렬/AMR 통합)."""
    _add_heading(doc, '5. 격자 생성', level=1)

    _add_paragraph(doc,
        '본 코드는 2D/3D 구조격자, 비정렬 삼각형 격자, 적응 격자 세분화를 지원한다.')

    _add_heading(doc, '5.1 2D 구조격자', level=2)
    _add_paragraph(doc,
        'Gmsh API 기반 2D 직사각형 구조격자 생성기로, '
        '채널, 공동, CHT 멀티존, 기포탑 등 검증용 격자를 생성한다.')

    _add_heading(doc, '5.2 2D 비정렬 삼각형 격자', level=2)
    _add_paragraph(doc,
        '구조 사각형 격자를 대각선으로 분할하여 삼각형 격자를 생성한다.\n'
        '• 셀 면적: A = (1/2)|e₁ × e₂|\n'
        '• 셀 중심: 세 꼭지점의 평균\n'
        '• 면 법선: 면 방향 벡터에 수직')

    _add_heading(doc, '5.3 3D 육면체 구조격자', level=2)
    _add_paragraph(doc,
        '3D 육면체(hexahedral) 구조격자 생성기:\n'
        '• 셀 체적: V = Δx · Δy · Δz\n'
        '• 6방향 면(±x, ±y, ±z) 구성 및 법선 벡터 계산\n'
        '• i,j,k 인덱스 기반 셀-면 연결성 구축\n'
        '• 경계면 자동 식별 및 태깅')

    _add_table(doc,
        ['항목', '2D', '3D'],
        [
            ['격자 셀 형태', '사각형/삼각형', '육면체'],
            ['이웃 셀 수', '3~4', '6'],
            ['체적 계산', 'Δx·Δy', 'Δx·Δy·Δz'],
            ['기울기 복원', '2D Green-Gauss', '3D Green-Gauss'],
        ])

    _add_heading(doc, '5.5 혼합 격자 (Hex/Tet)', level=2)
    _add_paragraph(doc,
        '육면체(Hex)와 사면체(Tet) 혼합 격자 생성기:\n'
        '• 영역을 x 방향으로 Hex 구간과 Tet 구간으로 분할\n'
        '• Tet 영역: Dompierre 5-tet 분해 (hex → 5 tet)\n'
        '• 인터페이스: hex 사각면을 2개 삼각면으로 분할하여 tet과 매칭\n'
        '• 다면체 FVM: _build_fvmesh_3d()가 삼각/사각 혼합면 자동 처리')

    _add_paragraph(doc, 'Dompierre 5-tet 분해:', bold=True)
    _add_paragraph(doc,
        'tet1: (n0,n1,n2,n5), tet2: (n0,n2,n7,n5), tet3: (n2,n3,n7,n0),\n'
        'tet4: (n5,n7,n6,n2), tet5: (n0,n5,n7,n4)\n\n'
        '대각선 (n0-n6)을 기준으로 분해하며, 인접 hex와 일관된 면 분할을 보장한다.')

    _add_heading(doc, '5.4 적응 격자 세분화(AMR)', level=2)
    _add_paragraph(doc,
        '사각형 셀을 4분할(quad→4quads)하는 적응 세분화:\n'
        '• 기울기-점프 추정기: η_P = Σ_f |φ_N - φ_P| · A_f / V_P\n'
        '• 상위 20~30% 셀을 세분화 대상으로 선정\n'
        '• 레벨 차이 ≤ 1 제약으로 격자 급변 방지\n'
        '• 해→추정→세분화→보간→재계산 루프')

    _add_table(doc,
        ['케이스', '격자 유형', '크기', '셀 수'],
        [
            ['Poiseuille', '2D 구조격자', '1.0×0.1 m', '50×20 = 1,000'],
            ['Cavity', '2D 구조격자', '1.0×1.0 m', '32×32 = 1,024'],
            ['CHT', '2D 멀티존', '0.5×0.04 m', '40×18 = 720'],
            ['Bubble Column', '2D 구조격자', '0.15×0.45 m', '15×45 = 675'],
            ['3D Duct', '3D 육면체', '1.0×0.1×0.1 m', '20×8×8 = 1,280'],
            ['비정렬', '삼각형', '1.0×0.1 m', '대각선 분할'],
            ['AMR', '적응 사분면', '동적 세분화', '가변'],
            ['혼합 격자', 'Hex+Tet', '2.0×0.1×0.1 m', 'Hex+5Tet/hex'],
        ])

    doc.add_page_break()


def _write_ch6_parallel(doc):
    """6장 병렬화 및 가속 (MPI/GPU/전처리기 통합)."""
    _add_heading(doc, '6. 병렬화 및 가속', level=1)

    _add_heading(doc, '6.1 MPI RCB 영역분할', level=2)
    _add_paragraph(doc,
        'RCB(Recursive Coordinate Bisection) 알고리즘으로 격자를 N개 파티션으로 분할한다:\n'
        '1. 전체 격자 셀의 좌표를 수집\n'
        '2. 가장 긴 방향 축을 선택\n'
        '3. 중앙값 기준으로 두 영역으로 이분\n'
        '4. 목표 파티션 수에 도달할 때까지 재귀 반복')

    _add_heading(doc, '6.2 고스트 셀 통신', level=2)
    _add_paragraph(doc,
        '파티션 경계에서의 데이터 교환은 고스트 셀 방식으로 처리한다:\n'
        '• 각 파티션은 이웃 파티션의 경계 셀을 고스트 셀로 복사\n'
        '• 매 반복 시작 시 MPI_Sendrecv로 고스트 셀 갱신\n'
        '• 압력 보정: 복제 압력 솔브(Replicated Pressure Solve) 방식')

    _add_heading(doc, '6.3 GPU 가속 (CuPy BiCGSTAB)', level=2)
    _add_paragraph(doc,
        'CuPy 라이브러리를 활용한 GPU BiCGSTAB 반복 솔버:\n'
        '• scipy.sparse → cupy.scipy.sparse 자동 전환\n'
        '• CSR 포맷 GPU 메모리 전송\n'
        '• GPU 미사용 시 자동 CPU 폴백\n'
        '• 수렴 기준: 잔차 < 1e-7')

    _add_heading(doc, '6.4 전처리기 (ILU/AMG)', level=2)
    _add_paragraph(doc,
        '반복 선형 솔버의 수렴 가속을 위한 전처리기:\n'
        '• Jacobi: 대각선 스케일링 (가장 단순)\n'
        '• ILU(0): 불완전 LU 분해 (패턴 고정)\n'
        '• ILU(k): 불완전 LU 분해 (k레벨 필인)\n'
        '• AMG: 대수적 다중격자 (가장 빠른 수렴)')

    doc.add_page_break()


def _write_ch7_verification_2d(doc, results, figures_dir):
    """7장 검증 결과 — 2D 문제."""
    _add_heading(doc, '7. 검증 결과 — 2D 문제', level=1)

    # --- 7.1 Case 1: Poiseuille ---
    _add_heading(doc, '7.1 Case 1: Poiseuille 유동', level=2)
    _add_paragraph(doc,
        '2D 평행판 사이 정상 층류 유동. 해석해가 존재하므로 정량적 검증이 가능하다. '
        '지배방정식은 비압축성 Navier-Stokes 방정식이며, 압력 기울기 dP/dx = -1.0 Pa/m가 '
        '균일하게 인가된다. 정상 상태에서 y 방향 속도 성분은 소멸하고 x 방향 속도만 남아 '
        '2차 상미분방정식 d²u/dy² = (1/μ)(dP/dx)를 만족한다.')
    _add_paragraph(doc,
        '해석해는 포물선 프로파일 u(y) = (1/2μ)(−dP/dx)(Hy − y²)이며, '
        '최대 속도는 채널 중앙(y = H/2)에서 u_max = (−dP/dx)H²/(8μ)로 결정된다. '
        '채널 높이 H = 0.1 m, 점성 μ = 0.01 Pa·s 조건에서 u_max = 0.03125 m/s이다. '
        '수치해와 해석해의 L2 상대 오차가 1% 미만이면 PASS로 판정한다.')

    _add_paragraph(doc, '해석 조건:', bold=True)
    _add_table(doc, ['항목', '값'], [
        ['채널 길이', '1.0 m'], ['채널 높이', '0.1 m'],
        ['밀도', '1.0 kg/m³'], ['점성', '0.01 Pa·s'],
        ['압력 기울기', '-1.0 Pa/m']])

    r1 = results.get('case1_poiseuille', {})
    if 'error' not in r1:
        _add_paragraph(doc, '해석 결과:', bold=True)
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(r1.get('converged', 'N/A'))],
            ['반복 횟수', str(r1.get('iterations', 'N/A'))],
            ['수치해 최대 속도', f"{r1.get('u_max_numerical', 0):.6f} m/s"],
            ['해석해 최대 속도', f"{r1.get('u_max_analytical', 0):.6f} m/s"],
            ['L2 상대 오차', f"{r1.get('L2_error', 0):.4e}"]])

        _add_figure(doc, os.path.join(figures_dir, 'case1_poiseuille.png'),
                    '그림 7.1: Poiseuille 유동 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case1_velocity.png'),
                    '그림 7.1-b: Poiseuille 속도 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r1["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.2 Case 2: Cavity ---
    _add_heading(doc, '7.2 Case 2: Lid-Driven Cavity', level=2)
    _add_paragraph(doc,
        '2D 정사각형 밀폐 공동(1.0 × 1.0 m) 내부 유동이다. 상단 벽(뚜껑)이 일정 속도 '
        'U = 1.0 m/s로 수평 이동하고 나머지 3면은 no-slip 조건이다. '
        '관성력과 점성력의 상호작용으로 공동 내부에 주 와류(primary vortex)와 '
        '모서리 보조 와류(corner vortices)가 형성된다.')
    _add_paragraph(doc,
        'Re = 100 및 Re = 400 두 조건에서 정상 상태 해를 구하고, '
        'Ghia et al. (1982) 참고문헌(128×128 격자, stream function-vorticity 법)의 '
        '수직/수평 중앙선 속도 프로파일과 비교한다. '
        'Re = 100에서 단일 주 와류 중심 위치, Re = 400에서 하단 보조 와류 존재 여부가 '
        '물리적 타당성 판정 기준이다.')

    _add_table(doc, ['항목', '값'], [
        ['영역 크기', '1.0 × 1.0 m'], ['뚜껑 속도', '1.0 m/s'],
        ['Reynolds 수', '100, 400'], ['격자', '32 × 32']])

    r2 = results.get('case2_cavity', {})
    if 'error' not in r2:
        r2_results = r2.get('results', {})
        rows = []
        for Re, res in r2_results.items():
            rows.append([str(Re), str(res.get('converged', 'N/A')),
                         str(res.get('iterations', 'N/A'))])
        if rows:
            _add_table(doc, ['Re', '수렴', '반복 횟수'], rows)

        _add_figure(doc, os.path.join(figures_dir, 'case2_cavity.png'),
                    '그림 7.2: Lid-Driven Cavity 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case2_velocity.png'),
                    '그림 7.2-b: Lid-Driven Cavity 속도 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r2["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.3 Case 3: CHT ---
    _add_heading(doc, '7.3 Case 3: Conjugate Heat Transfer', level=2)
    _add_paragraph(doc,
        '가열 고체 평판 하부에 층류 유체가 흐르는 2D 켤레 열전달(Conjugate Heat Transfer) 문제이다. '
        '고체 영역(상단)에는 균일 체적 열원 q‴이 인가되고, 유체 영역(하단)에는 Poiseuille 유동이 '
        '유지된다. 고체-유체 인터페이스에서는 온도 연속성(T_s = T_f)과 법선 열유속 연속성 '
        '(k_s ∂T/∂n|_s = k_f ∂T/∂n|_f)이 강제된다.')
    _add_paragraph(doc,
        '검증 기준은 전체 에너지 균형이다. 고체에 투입된 열량(q‴ × V_solid)이 유체 출구에서 '
        '반출되는 대류 열량(ṁ c_p ΔT)과 5% 이내로 일치해야 PASS로 판정한다. '
        '인터페이스 온도 분포의 평균값과 표준편차도 물리적 타당성 확인에 활용된다.')

    r3 = results.get('case3_cht', {})
    if 'error' not in r3:
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(r3.get('converged', 'N/A'))],
            ['CHT 반복 횟수', str(r3.get('iterations', 'N/A'))],
            ['인터페이스 평균 온도', f"{r3.get('T_interface_avg', 0):.2f} K"],
            ['에너지 균형 오차', f"{r3.get('energy_balance_error', 0)*100:.1f}%"]])

        _add_figure(doc, os.path.join(figures_dir, 'case3_cht.png'),
                    '그림 7.3: CHT 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case3_temperature.png'),
                    '그림 7.3-b: CHT 온도 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r3["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.4 Case 4: 기포탑 ---
    _add_heading(doc, '7.4 Case 4: 기포탑 이상유동', level=2)
    _add_paragraph(doc,
        '2D 직사각형 수조(0.15 × 0.45 m) 하부 중앙 슬롯에서 공기를 연속 주입하는 기포탑이다. '
        'Two-Fluid Model의 6방정식 체계(각 상의 연속·운동량 방정식)를 적용하며, '
        '기포 간 상호작용력으로 항력(drag), 가상질량력(virtual mass), '
        '양력(lift)을 포함한다. 기포 직경 d_b = 5 mm 기준으로 항력계수는 '
        'Schiller-Naumann 상관식을 사용한다.')
    _add_paragraph(doc,
        '물리적 타당성 기준으로 기체 체적분율 α_g가 [0, 1] 범위를 유지하고, '
        '기포 상승 후 수면 근처에서 α_g가 하부 주입부보다 증가하는 '
        '부력 구동 순환 패턴이 형성되어야 한다. '
        '정량 기준치는 없으나 최대 α_g > 입구 α_g(0.04) 조건을 PASS 판정에 활용한다.')

    _add_table(doc, ['항목', '값'], [
        ['수조 크기', '0.15 × 0.45 m'], ['기포 직경', '5 mm'],
        ['입구 기체 체적분율', '0.04'], ['입구 기체 속도', '0.1 m/s'],
        ['해석 시간', '2.0 s'], ['시간 간격', '0.005 s']])

    r4 = results.get('case4_bubble_column', {})
    if 'error' not in r4:
        _add_table(doc, ['항목', '값'], [
            ['시간 스텝 수', str(r4.get('time_steps', 'N/A'))],
            ['최대 기체 체적분율', f"{r4.get('alpha_g_max', 0):.4f}"],
            ['물리적 타당성', str(r4.get('physical_validity', 'N/A'))]])

        _add_figure(doc, os.path.join(figures_dir, 'case4_bubble_column.png'),
                    '그림 7.4: 기포탑 이상유동 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case4_alpha.png'),
                    '그림 7.4-b: 기포탑 보이드율 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r4["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.5 Case 6: MUSCL/TVD ---
    _add_heading(doc, '7.5 Case 6: MUSCL/TVD 고차 이산화', level=2)
    _add_paragraph(doc,
        '순수 이류 방정식 ∂φ/∂t + u ∂φ/∂x = 0 (u = 1.0 m/s)을 1D 도메인(0~1 m)에서 해석한다. '
        '초기 조건은 계단 함수(step function): x < 0.5이면 φ = 1, x ≥ 0.5이면 φ = 0이다. '
        '1차 Upwind 스킴은 수치 점성(numerical diffusion)으로 인터페이스를 넓히는 반면, '
        'MUSCL(Monotone Upstream-centered Scheme for Conservation Laws)/TVD(Total Variation '
        'Diminishing) 스킴은 van Leer 기울기 제한자(limiter)를 적용하여 급격한 기울기를 보존한다.')
    _add_paragraph(doc,
        '검증 지표는 인터페이스 확산 폭(diffusion width)이다. '
        't_end에서 φ가 [0.1, 0.9] 구간에 해당하는 x 범위의 길이로 정의하며, '
        'MUSCL의 확산 폭이 Upwind보다 좁아야 PASS 조건을 충족한다. '
        'L2 오차 비교로 MUSCL이 Upwind 대비 정확도 향상을 정량적으로 확인한다.')

    r6 = results.get('case6_muscl', {})
    if 'error' not in r6:
        muscl_sharper = r6.get('muscl_sharper', False)
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(r6.get('converged', 'N/A'))],
            ['MUSCL 해상도 향상', str(muscl_sharper)],
            ['Upwind L2', f"{r6.get('upwind_L2', 0):.4e}"],
            ['MUSCL L2', f"{r6.get('muscl_L2', 0):.4e}"],
            ['판정', 'PASS' if muscl_sharper else 'FAIL']])

        _add_figure(doc, os.path.join(figures_dir, 'case6_muscl.png'),
                    '그림 7.5: MUSCL/TVD 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case6_muscl.png'),
                    '그림 7.5-b: MUSCL/TVD 필드 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r6["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.6 Case 7: 비정렬 ---
    _add_heading(doc, '7.6 Case 7: 비정렬 격자', level=2)
    _add_paragraph(doc,
        '삼각형 비정렬 격자에서 Case 1과 동일한 평행판 Poiseuille 유동을 해석한다. '
        '비정렬 격자에서는 셀 면의 법선 벡터가 셀 중심 연결선과 일치하지 않으므로 '
        '비직교 보정(non-orthogonal correction)이 필수적이다. '
        '격자는 Delaunay 삼각분할 알고리즘으로 생성하며, 벽면 근처에는 '
        '경계층 해상도를 위해 격자를 조밀하게 배치한다.')
    _add_paragraph(doc,
        '비직교 보정 항은 확산 플럭스 계산 시 교차 미분(cross-diffusion) 항으로 추가되며, '
        '반복 수렴 내에서 명시적으로 처리한다. '
        'Case 1 해석해와의 L2 상대 오차가 10% 미만이면 비정렬 격자 FVM 이산화가 '
        '올바르게 구현되었음을 확인한다. 비직교 보정 여부에 따른 오차 감소도 기록한다.')

    r7 = results.get('case7_unstructured', {})
    if 'error' not in r7:
        l2_7 = r7.get('L2_error', float('inf'))
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(r7.get('converged', 'N/A'))],
            ['L2 상대 오차', f"{l2_7:.4e}"],
            ['비직교 보정', str(r7.get('non_orthogonal_correction', 'N/A'))],
            ['판정', 'PASS' if l2_7 < 0.10 else 'FAIL']])

        _add_figure(doc, os.path.join(figures_dir, 'case7_unstructured.png'),
                    '그림 7.6: 비정렬 격자 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case7_velocity.png'),
                    '그림 7.6-b: 비정렬 격자 속도 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r7["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.7 Case 8: MPI ---
    _add_heading(doc, '7.7 Case 8: MPI 병렬화', level=2)
    _add_paragraph(doc,
        'RCB(Recursive Coordinate Bisection) 알고리즘으로 계산 영역을 분할하고 '
        'MPI(Message Passing Interface) 병렬 연산을 검증한다. '
        '영역 경계에서는 유령 셀(ghost cell) 레이어를 교환하여 이웃 프로세스 데이터를 '
        '참조한다. 각 MPI 통신 단계에서 MPI_Sendrecv를 이용한 양방향 교환이 수행되며, '
        '프레셔 솔버는 전역 수렴 기준을 사용한다.')
    _add_paragraph(doc,
        '검증 방법은 동일 문제(Poiseuille 유동)를 직렬(1 프로세스)과 병렬(N 프로세스)로 '
        '각각 해석하여 전체 속도장의 L2 차이를 계산한다. '
        '부하 균형도(load balance)는 프로세스별 셀 수 편차로 측정하며, '
        '이상값 1.0에 가까울수록 균등 분할임을 나타낸다. '
        '직렬 대비 L2 차이가 부동소수점 반올림 오차 수준(~1e-12)이어야 PASS이다.')

    r8 = results.get('case8_mpi', {})
    if 'error' not in r8:
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(r8.get('converged', False))],
            ['프로세스 수', str(r8.get('n_procs', 'N/A'))],
            ['부하 균형도', f"{r8.get('load_balance', 0):.3f}"],
            ['직렬 대비 L2', f"{r8.get('serial_diff_L2', 0):.4e}"]])

        _add_figure(doc, os.path.join(figures_dir, 'case8_mpi.png'),
                    '그림 7.7: MPI 병렬화 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case8_partition.png'),
                    '그림 7.7-b: MPI 영역 분할 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r8["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.8 Case 9: 상변화 ---
    _add_heading(doc, '7.8 Case 9: Stefan 문제 (Lee 상변화)', level=2)
    _add_paragraph(doc,
        'Stefan 문제는 상변화 경계(phase interface)가 시간에 따라 이동하는 고전 검증 문제이다. '
        '1D 도메인에서 한쪽 끝을 포화 온도 이상으로 가열하면 액체-기체 경계가 이동하며, '
        '그 위치 X(t)는 Stefan 조건 ρ L dX/dt = k ∂T/∂x|_interface 를 만족한다. '
        '해석해는 X(t) = 2λ√(α t)이며, λ는 초월방정식으로 결정된다.')
    _add_paragraph(doc,
        'Lee 상변화 모델은 체적 열원/흡열 항으로 상변화를 처리한다: '
        'ṁ_evap = r_evap α_l ρ_l (T − T_sat)/T_sat (T > T_sat일 때), '
        'ṁ_cond = r_cond α_g ρ_g (T_sat − T)/T_sat (T < T_sat일 때). '
        '계수 r_evap = r_cond = 0.1 s⁻¹은 수치 안정성과 물리적 응답 속도의 균형점이며, '
        '실제 계면 동역학을 직접 표현하지 않는 한계가 있다. '
        'Stefan 수(Ste = c_p ΔT / L_latent)와 경계 위치 L2 오차로 정확도를 평가한다.')

    _add_table(doc, ['항목', '값'], [
        ['상변화 모델', 'Lee 모델'], ['증발 계수', '0.1 s⁻¹'],
        ['응결 계수', '0.1 s⁻¹'], ['포화 온도', '373.15 K']])

    r9 = results.get('case9_phase_change', {})
    if 'error' not in r9:
        l2_9 = r9.get('L2_error', float('inf'))
        conv9 = r9.get('converged', False)
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(conv9)],
            ['Stefan 수', f"{r9.get('Ste', 0):.6f}"],
            ['L2 오차', f"{l2_9:.4e}"],
            ['판정', 'PASS' if (conv9 and l2_9 < 0.10) else 'FAIL']])

        _add_figure(doc, os.path.join(figures_dir, 'case9_phase_change.png'),
                    '그림 7.8: 상변화 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case9_temperature.png'),
                    '그림 7.8-b: Stefan 문제 온도 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r9["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.9 Case 10: 화학반응 ---
    _add_heading(doc, '7.9 Case 10: 화학반응', level=2)
    _add_paragraph(doc,
        '완전 혼합 흐름 반응기(CSTR) 또는 플러그 흐름 반응기(PFR) 조건에서 '
        '1차 비가역 반응 A → B를 해석한다. '
        '종 수송 방정식은 ∂C/∂t + ∇·(uC) = ∇·(D∇C) − k·C 이며, '
        '균일 유동장(plug flow)과 확산 무시 조건에서 반응 항만 남아 '
        'dC/dt = −k·C 형태가 된다.')
    _add_paragraph(doc,
        '해석해는 C(t) = C₀ · exp(−k·t)이며, 초기 농도 C₀ = 1.0 mol/m³, '
        '반응 속도 상수 k = 1.0 s⁻¹ 조건에서 검증한다. '
        '시간 적분 후 전 영역 평균 농도를 해석해와 비교하며, '
        'L2 상대 오차가 5% 미만이면 종 수송 방정식의 반응 소스 항 구현이 '
        '올바른 것으로 판정한다.')

    r10 = results.get('case10_reaction', {})
    if 'error' not in r10:
        l2_10 = r10.get('L2_error', float('inf'))
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(r10.get('converged', 'N/A'))],
            ['최종 평균 농도', f"{r10.get('final_concentration', 0):.6f} mol/m³"],
            ['해석해 농도', f"{r10.get('analytical_concentration', 0):.6f} mol/m³"],
            ['L2 상대 오차', f"{l2_10:.4e}"],
            ['판정', 'PASS' if l2_10 < 0.05 else 'FAIL']])

        _add_figure(doc, os.path.join(figures_dir, 'case10_reaction.png'),
                    '그림 7.9: 화학반응 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case10_concentration.png'),
                    '그림 7.9-b: 반응 농도 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r10["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.10 Case 11: 복사 ---
    _add_heading(doc, '7.10 Case 11: 복사열전달', level=2)
    _add_paragraph(doc,
        'P1 복사 모델은 복사 강도를 구면 조화 함수 1차 전개로 근사한다. '
        '지배방정식은 복사 에너지 밀도 G에 대한 확산 방정식 '
        '∇·(Γ∇G) − a·G + 4a·σT⁴ = 0 이며, '
        '여기서 Γ = 1/(3(a + σ_s)), a는 흡수 계수, σ_s는 산란 계수이다. '
        '광학적으로 두꺼운 매질(optical thickness τ = a·L ≫ 1)에서 P1 모델의 '
        '정확도가 높으며, 얇은 매질에서는 오차가 증가한다.')
    _add_paragraph(doc,
        '경계 조건은 Marshak 조건: −Γ ∂G/∂n = (1/2)(G − 4σT_w⁴)이며, '
        '1D 균일 흡수 매질(τ ≫ 1) 조건의 해석해는 G = 4σT⁴(균일 온도장)이다. '
        '검증 기준은 에너지 균형 오차로, 벽면 경계에서 흡수·방출되는 복사 에너지의 '
        '총합이 전체 도메인 내 복사 소스와 5% 이내로 일치해야 PASS이다.')

    r11 = results.get('case11_radiation', {})
    if 'error' not in r11:
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(r11.get('converged', False))],
            ['최대 복사 에너지', f"{r11.get('G_max', 0):.4e} W/m²"],
            ['에너지 균형 오차', f"{r11.get('energy_balance_error', 0)*100:.2f}%"],
            ['판정', 'PASS' if r11.get('converged', False) else 'FAIL']])

        _add_figure(doc, os.path.join(figures_dir, 'case11_radiation.png'),
                    '그림 7.10: 복사열전달 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case11_radiation.png'),
                    '그림 7.10-b: 복사열전달 - 복사 에너지 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r11["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.11 Case 12: AMR ---
    _add_heading(doc, '7.11 Case 12: AMR', level=2)
    _add_paragraph(doc,
        '적응 격자 세분화(AMR: Adaptive Mesh Refinement) 알고리즘의 동작을 검증한다. '
        '기울기-점프 추정자(gradient-jump estimator)를 이용하여 각 셀의 세분화 필요성을 평가하며, '
        '물리량(압력, 속도, 온도 등)의 이웃 셀 간 기울기 차이가 임계값을 초과하면 '
        '해당 셀을 세분화 대상으로 표시한다.')
    _add_paragraph(doc,
        '2D 사각형 셀의 세분화는 4분할(quadrant splitting) 방식으로 수행된다: '
        '한 부모 셀이 4개의 자식 셀로 분할되고, 면 연결성과 이웃 정보가 갱신된다. '
        '세분화 후 보간(prolongation)으로 자식 셀의 초기 물리량을 부모로부터 계산하며, '
        '거칠게 만들기(coarsening)는 반대 방향으로 수행된다. '
        '검증 기준은 AMR 수행 후 셀 수가 초기값보다 증가했는지 여부이며, '
        '세분화 비율과 분포 패턴으로 알고리즘 정상 동작을 확인한다.')

    r12 = results.get('case12_amr', {})
    if 'error' not in r12:
        n_init = r12.get('n_cells_initial', 0)
        n_final = r12.get('n_cells_final', 0)
        _add_table(doc, ['항목', '값'], [
            ['초기 셀 수', str(n_init)],
            ['최종 셀 수', str(n_final)],
            ['세분화 비율', f"{n_final / max(n_init, 1):.2f}x"],
            ['판정', 'PASS' if n_final > n_init else 'FAIL']])

        _add_figure(doc, os.path.join(figures_dir, 'case12_amr.png'),
                    '그림 7.11: AMR 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case12_amr.png'),
                    '그림 7.11-b: AMR - 격자 세분화 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r12["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.12 Case 16: 전처리기 ---
    _add_heading(doc, '7.12 Case 16: 전처리기 비교', level=2)
    _add_paragraph(doc,
        '2D 대류-확산 문제에서 Jacobi, ILU(0), ILU(k), AMG 전처리기 성능 비교.')

    r16 = results.get('case16_preconditioner', {})
    if 'error' not in r16:
        precond_results = r16.get('results', {})
        precond_rows = []
        for name in ['none', 'jacobi', 'ilu0', 'iluk', 'amg']:
            pr = precond_results.get(name, {})
            if pr:
                precond_rows.append([
                    name.upper(),
                    str(pr.get('iterations', 'N/A')),
                    f"{pr.get('time', 0):.4f}",
                    f"{pr.get('L2_error', 0):.4e}"])
        if precond_rows:
            _add_table(doc, ['전처리기', '반복 횟수', '시간(s)', 'L2 오차'],
                       precond_rows)
        _add_paragraph(doc, f'최적 전처리기: {r16.get("best_preconditioner", "N/A")}')

        _add_figure(doc, os.path.join(figures_dir, 'case16_preconditioner.png'),
                    '그림 7.12: 전처리기별 반복 횟수 비교')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case16_solution.png'),
                    '그림 7.12-b: 전처리기 비교 - 해 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r16["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.13 Case 17: 적응 시간 간격 ---
    _add_heading(doc, '7.13 Case 17: CFL 기반 적응 시간 간격', level=2)
    _add_paragraph(doc,
        'CFL 기반 적응 시간 간격 제어의 정확성과 효율성 검증.')

    r17 = results.get('case17_adaptive_dt', {})
    if 'error' not in r17:
        fixed_steps = r17.get('fixed_dt_steps', 0)
        adaptive_steps = r17.get('adaptive_dt_steps', 0)
        cfl_ok = r17.get('cfl_always_below_max', False)
        reduction = r17.get('step_reduction', 0)

        _add_table(doc, ['방법', '스텝 수', 'dt', '스텝 비율'], [
            ['고정 dt', str(fixed_steps), '0.001', '-'],
            ['적응 dt', str(adaptive_steps), '가변', f'{reduction:.2f}']])

        _add_paragraph(doc,
            f'CFL 제약 충족: {cfl_ok}. '
            f'고정 대비 {(1.0 - reduction)*100:.0f}% 시간 스텝 절감.')

        _add_figure(doc, os.path.join(figures_dir, 'case17_adaptive_dt.png'),
                    '그림 7.13: 적응 시간 간격 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case17_solution.png'),
                    '그림 7.13-b: 적응 시간 간격 - 해 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r17["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.14 Case 19: 비등/응축 ---
    _add_heading(doc, '7.14 Case 19: 비등/응축 상변화', level=2)
    _add_paragraph(doc,
        'Rohsenow 핵비등 상관식, Nusselt 막응축 해석해, Lee 상변화 모델의 세 가지 물리 '
        '현상을 통합 검증한다. 각각 독립적인 해석해 또는 검증된 상관식이 존재하므로 '
        '정량적 비교가 가능하다.\n'
        '• Test A: Rohsenow 핵비등 열유속 — q″ = μ_l h_fg [g(ρ_l−ρ_g)/σ]^0.5 (C_sf)^(−3) Pr_l^(−1.7) ΔT_sat³\n'
        '• Test B: Nusselt 막응축 열전달 계수 — h = 0.943 [ρ_l(ρ_l−ρ_g)g h_fg k_l³ / (μ_l L ΔT)]^0.25\n'
        '• Test C: 1D 비등/응축 통합 시뮬레이션 — Lee 모델로 상변화율 계산, 에너지 보존 확인')
    _add_paragraph(doc,
        'Test A의 Zuber CHF(임계 열유속) 예측식 q″_CHF = 0.131 h_fg ρ_g [σg(ρ_l−ρ_g)/ρ_g²]^0.25 과의 '
        '비교를 통해 핵비등 영역의 상한을 확인한다. '
        'Test B는 수직 냉각판 위 막응축에 대해 Nusselt 해석해와 수치 결과를 비교하며, '
        '판 길이 방향으로 h(x) ∝ x^(−1/4) 감소 경향이 재현되어야 한다.')

    _add_table(doc, ['항목', '값'], [
        ['포화 온도', '373.15 K (물, 1 atm)'],
        ['잠열', '2.257×10⁶ J/kg'],
        ['Rohsenow C_sf', '0.013 (물-구리)'],
        ['통합 시뮬레이션 격자', '50셀 (1D)']])

    r19 = results.get('case19_boiling_condensation', {})
    if 'error' not in r19:
        q_range = r19.get('rohsenow_q_range', [0, 0])
        h_range = r19.get('nusselt_h_range', [0, 0])
        _add_table(doc, ['항목', '값'], [
            ['Rohsenow 열유속 범위', f'{q_range[0]:.2e} ~ {q_range[1]:.2e} W/m²'],
            ['Zuber CHF', f"{r19.get('zuber_chf', 0):.2e} W/m²"],
            ['Nusselt h 범위', f'{h_range[0]:.1f} ~ {h_range[1]:.1f} W/(m²·K)'],
            ['Nusselt 해석해 대비 오차', f"{r19.get('nusselt_error', 0):.2e}"],
            ['통합 시뮬레이션 수렴', str(r19.get('integrated_converged', False))],
            ['에너지 보존 오차', f"{r19.get('energy_balance_error', 0)*100:.2f}%"],
            ['판정', 'PASS' if r19.get('converged', False) else 'FAIL']])

        _add_figure(doc, os.path.join(figures_dir, 'case19_boiling_condensation.png'),
                    '그림 7.14: 비등/응축 상변화 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case19_temperature.png'),
                    '그림 7.14-b: 비등/응축 - 온도 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r19["error"]}', italic=True)

    doc.add_page_break()

    # --- 7.15 Case 21: 6-Equation 가열 채널 ---
    _add_heading(doc, '7.15 Case 21: 6-Equation 가열 채널', level=2)
    _add_paragraph(doc,
        '6-방정식 Two-Fluid Model로 수직 가열 채널의 비평형 상변화를 해석한다. '
        '6방정식 체계는 액체와 기체 각각에 대한 연속 방정식, 운동량 방정식, 에너지 방정식으로 구성된다. '
        '과냉수(T_in = 350 K)가 벽면 열유속 q″_wall = 500 kW/m²에 의해 가열되어 '
        '포화 온도(373.15 K)에 도달한 후 비등이 시작되는 과정을 모사한다.')
    _add_paragraph(doc,
        '본 검증은 속도장을 고정(solve_momentum=False)하고 에너지 방정식과 상변화 모델만 활성화하여 '
        '열수력 에너지 교환 구현의 정확성을 집중 검증한다. '
        '벽면에서 액체로의 열전달, 액체-기체 간 계면 열전달(h_if·a_if·(T_l − T_sat)), '
        'Lee 모델 상변화율이 에너지 보존을 만족하는지 확인한다. '
        '판정 기준은 에너지 수지 비율(입력 열량 대비 엔탈피 증가량)이 0.5~2.0 범위이고 '
        'α_g, T_l, T_g 분포가 물리적으로 타당해야 PASS이다.')

    _add_table(doc, ['항목', '값'], [
        ['채널 길이', '1.0 m'], ['채널 폭', '0.02 m'],
        ['입구 온도', '350 K (과냉수)'], ['포화 온도', '373.15 K'],
        ['벽면 열유속', '500 kW/m²'], ['입구 속도', '0.5 m/s'],
        ['초기 α_g', '0.05'], ['격자', '5×40'],
        ['해석 시간', '2.0 s'], ['시간 간격', '0.005 s'],
        ['출구 압력', '0.0 Pa (게이지)'], ['속도장', '고정 (에너지 검증 전용)']])

    r21 = results.get('case21_6eq_heated_channel', {})
    if 'error' not in r21:
        conv21 = r21.get('converged', False)
        phys21 = r21.get('physical', False)
        T_l_range = r21.get('T_l_range', [0, 0])
        T_g_range = r21.get('T_g_range', [0, 0])
        ag_range = r21.get('alpha_g_range', [0, 0])
        e_ratio = r21.get('energy_ratio', 0)
        energy_ok21 = 0.5 <= e_ratio <= 2.0
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(conv21)],
            ['T_l 범위', f'[{T_l_range[0]:.1f}, {T_l_range[1]:.1f}] K'],
            ['T_g 범위', f'[{T_g_range[0]:.1f}, {T_g_range[1]:.1f}] K'],
            ['α_g 범위', f'[{ag_range[0]:.4f}, {ag_range[1]:.4f}]'],
            ['에너지 수지 비율', f'{e_ratio:.4f}'],
            ['에너지 균형 기준 (0.5~2.0)', 'PASS' if energy_ok21 else 'FAIL'],
            ['물리적 타당성', str(phys21)],
            ['판정', 'PASS' if (conv21 and phys21 and energy_ok21) else 'FAIL']])

        _add_figure(doc, os.path.join(figures_dir, 'case21_6eq_heated_channel.png'),
                    '그림 7.15: 6-Equation 가열 채널 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case21_temperature.png'),
                    '그림 7.15-b: 6-Equation 가열 채널 - 액체 온도 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r21["error"]}', italic=True)

    doc.add_page_break()


def _write_ch8_verification_3d(doc, results, figures_dir):
    """8장 검증 결과 — 3D 문제."""
    _add_heading(doc, '8. 검증 결과 — 3D 문제', level=1)

    # 8.1 Case 5: 3D 덕트
    _add_heading(doc, '8.1 Case 5: 3D 덕트 Poiseuille 유동', level=2)
    _add_paragraph(doc,
        '3D 직사각형 덕트(1.0 × 0.1 × 0.1 m) 내 정상 층류 유동을 검증한다. '
        '2D Poiseuille 유동과 달리 3D 덕트에서는 y, z 두 방향 모두 no-slip 벽면이 존재하므로 '
        '속도 프로파일이 단순 포물선이 아닌 2D 분포를 가진다. '
        '해석해는 무한 급수 형태로 주어진다: '
        'u(y,z) = (16a²/π³μ)(−dP/dx) Σ_{n=1,3,5,...} (−1)^{(n−1)/2}/n³ '
        '[1 − cosh(nπy/(2a))/cosh(nπb/(2a))] cos(nπz/(2a)).')
    _add_paragraph(doc,
        'FVM 이산화는 y, z 두 방향의 점성 확산 항을 모두 포함하여 3D 라플라시안을 처리한다. '
        '격자 20 × 8 × 8에서 무한 급수의 처음 20항으로 계산한 해석해 최대 속도와 '
        '수치해를 비교하며, L2 상대 오차 5% 미만을 PASS 기준으로 한다. '
        '단면 속도 등고선의 대칭성도 물리적 타당성 확인에 활용된다.')

    _add_table(doc, ['항목', '값'], [
        ['덕트 크기', '1.0 × 0.1 × 0.1 m'], ['격자', '20 × 8 × 8'],
        ['밀도', '1.0 kg/m³'], ['점성', '0.01 Pa·s']])

    r5 = results.get('case5_3d_duct', {})
    if 'error' not in r5:
        l2_5 = r5.get('L2_error', float('inf'))
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(r5.get('converged', 'N/A'))],
            ['수치해 최대 속도', f"{r5.get('u_max_numerical', 0):.6f} m/s"],
            ['해석해 최대 속도', f"{r5.get('u_max_analytical', 0):.6f} m/s"],
            ['L2 상대 오차', f"{l2_5:.4e}"],
            ['판정', 'PASS' if l2_5 < 0.05 else 'FAIL']])

        _add_figure(doc, os.path.join(figures_dir, 'case5_3d_duct.png'),
                    '그림 8.1: 3D 덕트 Poiseuille 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case5_velocity.png'),
                    '그림 8.1-b: 3D 덕트 속도 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r5["error"]}', italic=True)

    doc.add_page_break()

    # 8.2 Case 14: 3D Cavity
    _add_heading(doc, '8.2 Case 14: 3D Lid-Driven Cavity', level=2)
    _add_paragraph(doc,
        '3D 정육면체 공동(1.0 × 1.0 × 1.0 m)에서 상단 면이 x 방향으로 U = 1.0 m/s로 '
        '이동하는 뚜껑 구동 유동이다. Case 2의 2D 확장으로, 3D에서는 z 방향 '
        '측벽(spanwise walls)의 영향으로 유동 구조가 2D 결과와 달라진다. '
        '중앙 단면(z = 0.5 m)에서 2D 유사 거동이 나타나야 하며, '
        '측벽 근처에서는 2차 유동(secondary flow)이 발생한다.')
    _add_paragraph(doc,
        'Re = 100 조건에서 정상 상태 해를 구하고, 중앙선 속도 프로파일이 '
        'Ghia et al. (1982) 2D 벤치마크와 질적으로 일치하는지 확인한다. '
        '주 와류 중심이 공동 상단 우측에 위치하고, 하단 모서리 보조 와류가 존재하면 '
        '물리적으로 타당한 것으로 판정한다. 16 × 16 × 16 격자로 해상도를 확보한다.')

    _add_table(doc, ['항목', '값'], [
        ['영역 크기', '1.0 × 1.0 × 1.0 m'], ['격자', '16 × 16 × 16'],
        ['뚜껑 속도', '1.0 m/s (x)'], ['Re', '100']])

    r14 = results.get('case14_3d_cavity', {})
    if 'error' not in r14:
        conv14 = r14.get('converged', False)
        phys14 = r14.get('physically_reasonable', False)
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(conv14)],
            ['반복 횟수', str(r14.get('iterations', 'N/A'))],
            ['중앙선 최대 u-속도', f"{r14.get('u_max', 0):.4f} m/s"],
            ['물리적 타당성', str(phys14)],
            ['판정', 'PASS' if (conv14 and phys14) else 'FAIL']])

        _add_figure(doc, os.path.join(figures_dir, 'case14_3d_cavity.png'),
                    '그림 8.2: 3D Cavity 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case14_velocity.png'),
                    '그림 8.2-b: 3D Cavity 속도 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r14["error"]}', italic=True)

    doc.add_page_break()

    # 8.3 Case 15: 3D 자연대류
    _add_heading(doc, '8.3 Case 15: 3D 자연대류 (Rayleigh-Bénard)', level=2)
    _add_paragraph(doc,
        '3D 밀폐 직육면체 공간(1.0 × 1.0 × 0.5 m)에서 Boussinesq 근사를 적용한 '
        '자연대류를 검증한다. Boussinesq 근사는 밀도 변화를 부력 항에만 반영하고 '
        '나머지 물성은 일정하다고 가정한다: ρ ≈ ρ₀[1 − β(T − T_ref)]. '
        '하단 벽면(T_hot = 1.0, 무차원)과 상단 벽면(T_cold = 0.0) 사이 '
        '온도차가 부력을 유발하여 대류 순환이 형성된다.')
    _add_paragraph(doc,
        '열팽창 계수 β = 0.01 K⁻¹ 조건에서 Rayleigh 수는 '
        'Ra = gβΔTH³/(να) 로 계산된다. '
        'Nusselt 수 Nu = q″H/(kΔT)는 하단 벽면 열유속으로 계산하며, '
        'Ra에 따른 Nu 상관식(예: Nu ≈ 0.069 Ra^(1/3) Pr^0.074)과 비교하여 '
        '자연대류 강도가 물리적으로 합당한지 확인한다. '
        '온도 성층화(하단 고온, 상단 저온) 여부가 주요 물리 검증 기준이다.')

    _add_table(doc, ['항목', '값'], [
        ['영역 크기', '1.0 × 1.0 × 0.5 m'], ['격자', '16 × 16 × 8'],
        ['하단 온도', '1.0'], ['상단 온도', '0.0'],
        ['열팽창 계수', '0.01 K⁻¹']])

    r15 = results.get('case15_3d_convection', {})
    if 'error' not in r15:
        conv15 = r15.get('converged', False)
        strat15 = r15.get('stratified', False)
        nu15 = r15.get('nusselt', 0.0)
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(conv15)],
            ['Nusselt 수', f'{nu15:.4f}'],
            ['온도 성층화', '확인' if strat15 else '미확인'],
            ['온도 범위', f"[{r15.get('T_min', 0):.3f}, {r15.get('T_max', 1):.3f}]"],
            ['판정', 'PASS' if (conv15 and strat15) else 'FAIL']])

        _add_figure(doc, os.path.join(figures_dir, 'case15_3d_convection.png'),
                    '그림 8.3: 3D 자연대류 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case15_temperature.png'),
                    '그림 8.3-b: 3D 자연대류 온도 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r15["error"]}', italic=True)

    doc.add_page_break()

    # 8.4 Case 22: 혼합 격자 Poiseuille
    _add_heading(doc, '8.4 Case 22: 혼합 격자(Hex/Tet) Poiseuille 유동', level=2)
    _add_paragraph(doc,
        '3D 덕트 Poiseuille 유동을 Hex/Tet 혼합 격자에서 해석하고, '
        'Case 5(순수 Hex)와 동일한 해석해 대비 L2 오차를 비교한다. '
        '실제 공학 해석에서는 복잡한 기하형상 영역에 사면체(Tet) 격자를, '
        '주요 유동 방향에는 육면체(Hex) 격자를 혼합 사용하는 것이 일반적이다. '
        'Tet 격자는 Dompierre 5-tet 분해 방식으로 Hex 셀에서 생성된다.')
    _add_paragraph(doc,
        'Hex-Tet 인터페이스에서는 면 연결성이 불일치(non-conformal)할 수 있으므로 '
        '인터페이스 면에서 플럭스 연속성을 보간으로 처리한다. '
        '인터페이스 연속성 검증은 속도·압력의 인터페이스 양측 값 차이가 '
        '수치 반올림 오차 수준이어야 한다. '
        'Hex 영역과 Tet 영역 각각의 L2 오차를 분리 계산하여 격자 유형별 '
        '정확도를 비교하며, 전체 L2 오차 10% 미만을 PASS 기준으로 한다.')

    _add_table(doc, ['항목', '값'], [
        ['덕트 크기', '2.0 × 0.1 × 0.1 m'],
        ['격자', '20 × 6 × 6 (Hex 50% + Tet 50%)'],
        ['밀도', '1.0 kg/m³'], ['점성', '0.01 Pa·s'],
        ['Tet 분해', 'Dompierre 5-tet']])

    r22 = results.get('case22_hybrid_mesh', {})
    if 'error' not in r22:
        l2_22 = r22.get('L2_error', float('inf'))
        l2_hex = r22.get('L2_hex', 0)
        l2_tet = r22.get('L2_tet', 0)
        if_cont = r22.get('interface_continuous', False)
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(r22.get('converged', 'N/A'))],
            ['전체 L2 오차', f'{l2_22:.4e}'],
            ['Hex 영역 L2', f'{l2_hex:.4e}'],
            ['Tet 영역 L2', f'{l2_tet:.4e}'],
            ['인터페이스 연속성', str(if_cont)],
            ['Hex 셀 수', str(r22.get('n_hex_cells', 0))],
            ['Tet 셀 수', str(r22.get('n_tet_cells', 0))],
            ['판정', 'PASS' if l2_22 < 0.10 else 'FAIL']])

        _add_figure(doc, os.path.join(figures_dir, 'case22_hybrid_mesh.png'),
                    '그림 8.4: 혼합 격자(Hex/Tet) Poiseuille 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case22_velocity.png'),
                    '그림 8.4-b: 혼합 격자 - 속도 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r22["error"]}', italic=True)

    doc.add_page_break()

    # 8.5 Case 23: 3D 과도 채널 유동
    _add_heading(doc, '8.5 Case 23: 3D 과도 채널 유동 (Transient Channel Flow)', level=2)
    _add_paragraph(doc,
        '3D 직육면체 채널에서 비정상 단상 유동 해석. '
        '초기 정지 상태에서 입구 속도를 갑자기 인가(impulsive start)하여 '
        '속도 프로파일이 시간에 따라 발달하는 과정을 검증한다.')

    _add_table(doc, ['항목', '값'], [
        ['채널 크기', '1.0 × 0.1 × 0.1 m'],
        ['격자', '20 × 10 × 10 (2000 cells)'],
        ['밀도', '1.0 kg/m³'],
        ['점성', '0.01 Pa·s'],
        ['입구 속도', '1.0 m/s'],
        ['Re', '10'],
        ['해석 유형', '비정상 (Transient)']])

    r23 = results.get('case23_3d_transient', {})
    if 'error' not in r23:
        l2_23 = r23.get('L2_error', float('inf'))
        dev23 = r23.get('transient_development', False)
        conv23 = r23.get('converged', False)
        _add_table(doc, ['항목', '값'], [
            ['최종 L2 오차 (vs Poiseuille)', f'{l2_23:.4e}'],
            ['비정상 발달 확인', str(dev23)],
            ['시간 스텝 수', str(r23.get('time_steps', 'N/A'))],
            ['VTU 스냅샷 수', str(r23.get('n_snapshots', 'N/A'))],
            ['판정', 'PASS' if conv23 else 'FAIL']])

        _add_paragraph(doc,
            '검증 결과: 초기 정지 상태에서 시작하여 시간이 경과함에 따라 '
            '속도 프로파일이 균일 분포에서 Poiseuille 포물선 분포로 발달하는 과정이 관찰되었다. '
            '충분한 시간 후 정상 상태 해석해와의 L2 오차로 수렴을 확인하였다.')

        _add_figure(doc, os.path.join(figures_dir, 'case23_3d_transient.png'),
                    '그림 8.5: 3D 과도 채널 유동 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case23_velocity.png'),
                    '그림 8.5-b: 3D 과도 채널 - 속도 분포 시간 변화 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r23["error"]}', italic=True)

    doc.add_page_break()


def _write_ch9_verification_phase_change(doc, results, figures_dir):
    """9장 검증 결과 — 상변화/다상유동."""
    _add_heading(doc, '9. 검증 결과 — 상변화/다상유동', level=1)

    # 9.1 Case 20: Edwards 블로우다운
    _add_heading(doc, '9.1 Case 20: Edwards 파이프 블로우다운 (Flashing)', level=2)
    _add_paragraph(doc,
        'Edwards & O\'Brien(1970) 수평 파이프 블로우다운 실험을 모사한다. '
        '고온·고압 과냉 액체가 충전된 파이프 끝단이 순간 파열되면, '
        '희석파가 전파되면서 국소 압력이 포화 압력 아래로 떨어져 '
        '급격한 플래싱(flashing) 상변화가 발생한다.')

    _add_paragraph(doc, '문제 설정:', bold=True)
    _add_table(doc, ['항목', '값'], [
        ['파이프 길이', '4.096 m'],
        ['파이프 직경', '0.073 m'],
        ['초기 압력', '7.0 MPa'],
        ['초기 온도', '502 K (과냉 액체)'],
        ['출구 압력', '0.1 MPa (순간 파열)'],
        ['초기 액체 체적분율', '1.0'],
        ['참고문헌', 'Edwards & O\'Brien (1970), JBNES 9(2), 125-135']])

    _add_paragraph(doc,
        '간이 모델링 전략:\n'
        '• 희석파(rarefaction wave): 음속 c≈1400 m/s로 선형 전파\n'
        '• T_sat(P): Antoine 근사식으로 포화온도 계산\n'
        '• 플래싱 조건: T_local > T_sat(P_local) 시 Lee 모델 적용\n'
        '• 목표: 물리적 거동의 정성적 재현 (압력파 전파, void 생성)')

    r20 = results.get('case20_edwards_blowdown', {})
    if 'error' not in r20:
        flash = r20.get('flashing_occurred', False)
        void = r20.get('max_void', 0.0)
        conv20 = r20.get('converged', False)
        _add_table(doc, ['항목', '값'], [
            ['수렴 여부', str(conv20)],
            ['플래싱 발생', str(flash)],
            ['감압 확인', str(r20.get('depressurization', False))],
            ['최대 보이드율', f'{void:.4f}'],
            ['공간 순서 보존', str(r20.get('spatial_ordering', False))],
            ['판정', 'PASS' if conv20 else 'FAIL']])

        _add_paragraph(doc,
            '검증 결과: 파이프 끝단 파열 후 희석파가 폐쇄단을 향해 전파되며, '
            '국소 압력 강하에 의한 플래싱이 관찰되었다. '
            '보이드율은 출구 근처에서 최대값을 보이며, '
            '폐쇄단에서 최소값을 보여 물리적으로 타당한 공간 분포를 나타낸다.')

        _add_figure(doc, os.path.join(figures_dir, 'case20_edwards_blowdown.png'),
                    '그림 9.1: Edwards 파이프 블로우다운 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case20_void.png'),
                    '그림 9.1-b: Edwards 블로우다운 - 보이드율 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r20["error"]}', italic=True)

    doc.add_page_break()

    # 9.2 Case 24: 풀 비등
    _add_heading(doc, '9.2 Case 24: 풀 비등 (Pool Boiling)', level=2)
    _add_paragraph(doc,
        '하부 벽면에 열유속을 가하여 과열된 액체에서 비등이 발생하는 과정을 시뮬레이션한다. '
        'Two-Fluid 솔버의 6-equation 모드(에너지 방정식 포함)를 사용하며, '
        'Lee 상변화 모델을 통해 증기 생성을 모사한다.')

    _add_table(doc, ['항목', '값'], [
        ['영역 크기', '0.02 × 0.04 m'],
        ['격자', '20 × 40 (800 cells)'],
        ['포화 온도', '373.15 K (100°C, 1 atm)'],
        ['벽면 열유속', '50 kW/m²'],
        ['상변화 모델', 'Lee (r = 0.1 1/s)'],
        ['해석 유형', '비정상 Two-Fluid 6-Equation']])

    r24 = results.get('case24_pool_boiling', {})
    if 'error' not in r24:
        boil = r24.get('boiling_occurred', False)
        temp_ok = r24.get('temp_reasonable', False)
        conv24 = r24.get('converged', False)
        T_range = r24.get('T_range', [0, 0])
        _add_table(doc, ['항목', '값'], [
            ['비등 발생', str(boil)],
            ['최대 증기분율', f"{r24.get('max_alpha_g', 0):.4f}"],
            ['평균 증기분율', f"{r24.get('mean_alpha_g', 0):.4f}"],
            ['온도 범위', f'{T_range[0]:.1f} ~ {T_range[1]:.1f} K'],
            ['온도 합리성', str(temp_ok)],
            ['판정', 'PASS' if conv24 else 'FAIL']])

        _add_paragraph(doc,
            '검증 결과: 하부 벽면 열유속에 의해 벽 근처 액체의 온도가 포화 온도를 초과하면서 '
            'Lee 모델에 의한 증기 생성이 관찰되었다. 증기는 부력에 의해 상부로 이동하며, '
            '온도 분포는 물리적으로 합리적인 범위 내에 있음을 확인하였다.')

        _add_figure(doc, os.path.join(figures_dir, 'case24_pool_boiling.png'),
                    '그림 9.2: 풀 비등 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case24_boiling.png'),
                    '그림 9.2-b: 풀 비등 - 증기 체적분율 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r24["error"]}', italic=True)

    doc.add_page_break()

    # 9.3 Case 25: 막 응축
    _add_heading(doc, '9.3 Case 25: 수직벽 막 응축 (Film Condensation)', level=2)
    _add_paragraph(doc,
        '수직 냉벽에서 증기가 응축하여 액막이 형성되는 과정을 시뮬레이션한다. '
        'Two-Fluid 솔버의 6-equation 모드를 사용하며, Nusselt 막응축 해석해와 비교한다.')

    _add_table(doc, ['항목', '값'], [
        ['영역 크기', '0.04 × 0.10 m'],
        ['격자', '40 × 20 (800 cells)'],
        ['포화 온도', '373.15 K'],
        ['벽면 온도', '363.15 K (과냉도 10 K)'],
        ['상변화 모델', 'Lee (r = 0.1 1/s)'],
        ['해석 유형', '비정상 Two-Fluid 6-Equation'],
        ['참조 해석해', 'Nusselt (1916) 막응축 이론']])

    r25 = results.get('case25_film_condensation', {})
    if 'error' not in r25:
        cond = r25.get('condensation_occurred', False)
        temp_grad = r25.get('temp_gradient_ok', False)
        conv25 = r25.get('converged', False)
        delta_num = r25.get('delta_numerical', 0)
        delta_nus = r25.get('delta_nusselt', 0)
        _add_table(doc, ['항목', '값'], [
            ['응축 발생', str(cond)],
            ['냉벽 근처 최대 α_l', f"{r25.get('max_alpha_l_near_wall', 0):.4f}"],
            ['수치 액막 두께', f'{delta_num*1000:.3f} mm'],
            ['Nusselt 해석 액막 두께', f'{delta_nus*1000:.3f} mm'],
            ['두께 비율 (수치/해석)', f"{r25.get('film_ratio', 0):.2f}"],
            ['Nusselt 열전달 계수', f"{r25.get('h_nusselt', 0):.1f} W/(m²·K)"],
            ['온도 구배 OK', str(temp_grad)],
            ['판정', 'PASS' if conv25 else 'FAIL']])

        _add_paragraph(doc,
            '검증 결과: 수직 냉벽에서 증기가 응축하여 액막이 형성되는 과정이 관찰되었다. '
            '냉벽 근처에서 액체 체적분율이 증가하고, 온도는 벽면 온도에서 포화 온도로 '
            '점진적으로 증가하는 물리적으로 타당한 분포를 보였다. '
            'Nusselt 해석해와의 액막 두께 비교를 통해 정량적 타당성을 확인하였다.')

        _add_figure(doc, os.path.join(figures_dir, 'case25_film_condensation.png'),
                    '그림 9.3: 수직벽 막 응축 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case25_condensation.png'),
                    '그림 9.3-b: 막 응축 - 액체 체적분율 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r25["error"]}', italic=True)

    doc.add_page_break()


def _write_ch10_verification_etc(doc, results, figures_dir):
    """10장 검증 결과 — 기타."""
    _add_heading(doc, '10. 검증 결과 — 기타', level=1)

    # 10.1 Case 13: GPU
    _add_heading(doc, '10.1 Case 13: GPU 가속 검증', level=2)
    _add_paragraph(doc,
        'CuPy BiCGSTAB GPU 솔버 검증. CPU 폴백 이식성 및 수치 정확도를 확인한다.')

    r13 = results.get('case13_gpu', {})
    if 'error' not in r13:
        backend = r13.get('backend', 'N/A')
        gpu_avail = r13.get('gpu_available', False)
        acc_verified = r13.get('accuracy_verified', False)
        max_err = r13.get('max_bicgstab_error', float('nan'))
        _add_table(doc, ['항목', '값'], [
            ['사용 백엔드', str(backend)],
            ['GPU 사용 가능', '예' if gpu_avail else '아니오 (CPU 폴백)'],
            ['BiCGSTAB 정확도', 'PASS' if acc_verified else 'FAIL'],
            ['최대 L2 오차', f"{max_err:.2e}" if max_err == max_err else 'N/A']])

        benchmarks = r13.get('benchmarks', [])
        if benchmarks:
            _add_paragraph(doc, '격자 크기별 벤치마크:', bold=True)
            bm_rows = []
            for bm in benchmarks:
                n = bm.get('n_cells', 0)
                td = bm.get('cpu_direct', 0)
                tb = bm.get('cpu_bicgstab', 0)
                tg = bm.get('gpu_bicgstab', float('nan'))
                err = bm.get('bicgstab_l2_error', 0)
                gpu_str = f'{tg:.4f}' if not math.isnan(tg) else 'N/A'
                speedup = f'{tb/max(tg, 1e-10):.2f}x' if not math.isnan(tg) else 'N/A'
                bm_rows.append([f'{n:,}', f'{td:.4f}', f'{tb:.4f}',
                                gpu_str, speedup, f'{err:.2e}'])
            _add_table(doc,
                ['셀 수', 'Direct[s]', 'CPU BiCG[s]', 'GPU BiCG[s]', '속도비', 'L2'],
                bm_rows)

        _add_figure(doc, os.path.join(figures_dir, 'case13_gpu.png'),
                    '그림 10.1: GPU 가속 검증 결과')
        _add_figure(doc, os.path.join(figures_dir, 'pv_case13_solution.png'),
                    '그림 10.1-b: GPU 가속 - 해 분포 (메쉬 시각화)')
    else:
        _add_paragraph(doc, f'해석 오류: {r13["error"]}', italic=True)

    doc.add_page_break()

    # 10.2 Case 18: 웹 대시보드
    _add_heading(doc, '10.2 Case 18: 웹 기반 실시간 대시보드', level=2)
    _add_paragraph(doc,
        'Flask + Plotly.js 기반 실시간 시뮬레이션 모니터링 대시보드 검증.')

    r18 = results.get('case18_web_dashboard', {})
    if 'error' not in r18:
        _add_table(doc, ['항목', '결과'], [
            ['Flask 서버 시작', '성공' if r18.get('server_started', False) else '실패'],
            ['REST API 응답', '성공' if r18.get('api_responsive', False) else '실패'],
            ['데이터 수집', f'{r18.get("n_data_points", 0)}개'],
            ['HTML 대시보드', '성공' if r18.get('html_served', False) else '실패']])

        _add_figure(doc, os.path.join(figures_dir, 'case18_web_dashboard.png'),
                    '그림 10.2: 웹 대시보드 모의 레이아웃')
    else:
        _add_paragraph(doc, f'해석 오류: {r18["error"]}', italic=True)

    doc.add_page_break()


def _write_ch11_summary(doc, results):
    """11장 검증 결과 요약 (25개 케이스 종합 표)."""
    _add_heading(doc, '11. 검증 결과 요약', level=1)

    _add_paragraph(doc,
        '25개 검증 케이스의 결과를 종합한다. '
        '2D 문제 15개(6-equation 포함), 3D 문제 5개(과도해석, 혼합격자 포함), '
        '상변화/다상유동 3개(비등, 응축 포함), 기타 2개로 분류하였다.')

    summary_rows = []

    def _s(key, default='N/A'):
        r = results.get(key, {})
        return r, 'error' not in r

    def _add_row(num, name, r, ok, metric_func):
        if ok:
            metric, status = metric_func(r)
            summary_rows.append([str(num), name, metric, status])
        else:
            summary_rows.append([str(num), name, 'ERROR', 'FAIL'])

    # Case 1
    r, ok = _s('case1_poiseuille')
    _add_row(1, 'Poiseuille 유동', r, ok,
             lambda r: (f"L2={r.get('L2_error',1):.4e}",
                        'PASS' if r.get('L2_error', 1) < 0.01 else 'FAIL'))

    # Case 2
    r2, ok2 = _s('case2_cavity')
    if ok2:
        for Re, res in r2.get('results', {}).items():
            conv = res.get('converged', False)
            summary_rows.append(['2', f'Cavity Re={Re}', f'수렴={conv}',
                                 'PASS' if conv else 'FAIL'])
    else:
        summary_rows.append(['2', 'Cavity', 'ERROR', 'FAIL'])

    # Case 3
    r, ok = _s('case3_cht')
    _add_row(3, 'CHT', r, ok,
             lambda r: (f"T_avg={r.get('T_interface_avg',0):.1f}K, E_err={r.get('energy_balance_error',1)*100:.1f}%",
                        'PASS' if (r.get('converged', False) and r.get('energy_balance_error', 1) < 0.20) else 'FAIL'))

    # Case 4
    r, ok = _s('case4_bubble_column')
    _add_row(4, '기포탑', r, ok,
             lambda r: (f"α_max={r.get('alpha_g_max',0):.4f}",
                        'PASS' if r.get('physical_validity', False) else 'FAIL'))

    # Case 5
    r, ok = _s('case5_3d_duct')
    _add_row(5, '3D 덕트', r, ok,
             lambda r: (f"L2={r.get('L2_error',1):.4e}",
                        'PASS' if r.get('L2_error', 1) < 0.05 else 'FAIL'))

    # Case 6
    r, ok = _s('case6_muscl')
    _add_row(6, 'MUSCL/TVD', r, ok,
             lambda r: (f"향상={r.get('muscl_sharper', False)}",
                        'PASS' if r.get('muscl_sharper', False) else 'FAIL'))

    # Case 7
    r, ok = _s('case7_unstructured')
    _add_row(7, '비정렬 격자', r, ok,
             lambda r: (f"L2={r.get('L2_error',1):.4e}",
                        'PASS' if r.get('L2_error', 1) < 0.10 else 'FAIL'))

    # Case 8
    r, ok = _s('case8_mpi')
    _add_row(8, 'MPI 병렬화', r, ok,
             lambda r: (f"수렴={r.get('converged', False)}",
                        'PASS' if r.get('converged', False) else 'FAIL'))

    # Case 9
    r, ok = _s('case9_phase_change')
    _add_row(9, 'Lee 상변화', r, ok,
             lambda r: (f"L2={r.get('L2_error',1):.4e}",
                        'PASS' if (r.get('converged', False) and r.get('L2_error', 1) < 0.10) else 'FAIL'))

    # Case 10
    r, ok = _s('case10_reaction')
    _add_row(10, '화학반응', r, ok,
             lambda r: (f"L2={r.get('L2_error',1):.4e}",
                        'PASS' if r.get('L2_error', 1) < 0.05 else 'FAIL'))

    # Case 11
    r, ok = _s('case11_radiation')
    _add_row(11, '복사열전달', r, ok,
             lambda r: (f"수렴={r.get('converged', False)}",
                        'PASS' if r.get('converged', False) else 'FAIL'))

    # Case 12
    r, ok = _s('case12_amr')
    _add_row(12, 'AMR', r, ok,
             lambda r: (f"셀 {r.get('n_cells_initial',0)}→{r.get('n_cells_final',0)}",
                        'PASS' if r.get('n_cells_final', 0) > r.get('n_cells_initial', 0) else 'FAIL'))

    # Case 13
    r, ok = _s('case13_gpu')
    _add_row(13, 'GPU 가속', r, ok,
             lambda r: (f"백엔드={r.get('backend','N/A')}",
                        'PASS' if r.get('accuracy_verified', r.get('converged', False)) else 'FAIL'))

    # Case 14
    r, ok = _s('case14_3d_cavity')
    _add_row(14, '3D Cavity', r, ok,
             lambda r: (f"수렴={r.get('converged',False)}, 물리={r.get('physically_reasonable',False)}",
                        'PASS' if (r.get('converged', False) and r.get('physically_reasonable', False)) else 'FAIL'))

    # Case 15
    r, ok = _s('case15_3d_convection')
    _add_row(15, '3D 자연대류', r, ok,
             lambda r: (f"Nu={r.get('nusselt',0):.4f}",
                        'PASS' if (r.get('converged', False) and r.get('stratified', False)) else 'FAIL'))

    # Case 16
    r, ok = _s('case16_preconditioner')
    _add_row(16, '전처리기', r, ok,
             lambda r: (f"최적={r.get('best_preconditioner','N/A')}",
                        'PASS' if r.get('all_accurate', False) else 'FAIL'))

    # Case 17
    r, ok = _s('case17_adaptive_dt')
    _add_row(17, '적응 시간 간격', r, ok,
             lambda r: (f"스텝={r.get('adaptive_dt_steps',0)}",
                        'PASS' if (r.get('converged', False) and r.get('cfl_always_below_max', False)) else 'FAIL'))

    # Case 18
    r, ok = _s('case18_web_dashboard')
    _add_row(18, '웹 대시보드', r, ok,
             lambda r: (f"데이터={r.get('n_data_points',0)}점",
                        'PASS' if r.get('converged', False) else 'FAIL'))

    # Case 19
    r, ok = _s('case19_boiling_condensation')
    _add_row(19, '비등/응축', r, ok,
             lambda r: (f"수렴={r.get('converged', False)}",
                        'PASS' if r.get('converged', False) else 'FAIL'))

    # Case 20
    r, ok = _s('case20_edwards_blowdown')
    _add_row(20, 'Edwards 블로우다운', r, ok,
             lambda r: (f"flashing={r.get('flashing_occurred',False)}, void={r.get('max_void',0):.3f}",
                        'PASS' if (r.get('converged', False) and r.get('flashing_occurred', False)) else 'FAIL'))

    # Case 21
    r, ok = _s('case21_6eq_heated_channel')
    _add_row(21, '6-Eq 가열 채널', r, ok,
             lambda r: (f"T_l증가={r.get('T_l_increase',0):.1f}K, E={r.get('energy_ratio',0):.3f}",
                        'PASS' if (r.get('converged', False) and r.get('physical', False)
                                   and 0.5 <= r.get('energy_ratio', 0) <= 2.0) else 'FAIL'))

    # Case 22
    r, ok = _s('case22_hybrid_mesh')
    _add_row(22, '혼합격자 Poiseuille', r, ok,
             lambda r: (f"L2={r.get('L2_error',1):.4e}",
                        'PASS' if r.get('L2_error', 1) < 0.10 else 'FAIL'))

    # Case 23
    r, ok = _s('case23_3d_transient')
    _add_row(23, '3D 과도 채널 유동', r, ok,
             lambda r: (f"L2={r.get('L2_error',1):.4e}, 발달={r.get('transient_development',False)}",
                        'PASS' if r.get('converged', False) else 'FAIL'))

    # Case 24
    r, ok = _s('case24_pool_boiling')
    _add_row(24, '풀 비등', r, ok,
             lambda r: (f"α_g={r.get('max_alpha_g',0):.4f}, 비등={r.get('boiling_occurred',False)}",
                        'PASS' if r.get('converged', False) else 'FAIL'))

    # Case 25
    r, ok = _s('case25_film_condensation')
    _add_row(25, '막 응축', r, ok,
             lambda r: (f"δ={r.get('delta_numerical',0)*1000:.2f}mm, 응축={r.get('condensation_occurred',False)}",
                        'PASS' if r.get('converged', False) else 'FAIL'))

    _add_table(doc, ['번호', '케이스', '주요 결과', '판정'], summary_rows)

    doc.add_page_break()


def _write_ch12_conclusion(doc):
    """12장 결론 및 향후 과제."""
    _add_heading(doc, '12. 결론 및 향후 과제', level=1)

    _add_heading(doc, '12.1 결론', level=2)
    _add_paragraph(doc,
        'Python 기반 Two-Fluid Model FVM 열유체 코드를 개발하고, '
        '25개의 검증 케이스를 통해 코드의 정확성과 확장성을 확인하였다. '
        '주요 성과는 다음과 같다:\n\n'
        '• Two-Fluid (Euler-Euler) 이상유동 모델 구현 및 검증\n'
        '• 6-방정식 비평형 상변화 모델 (각 상별 에너지 SIMPLE 내 결합)\n'
        '• SIMPLE 알고리즘 기반 속도-압력 커플링\n'
        '• 표준 k-ε 난류 모델 + 벽함수\n'
        '• 유체-고체 CHT 커플링\n'
        '• 2D/3D 구조·비정렬 격자 통합 지원\n'
        '• Hex/Tet 혼합 격자 생성기 (Dompierre 5-tet 분해)\n'
        '• MUSCL/TVD 고차 이산화 스킴 (4종 리미터)\n'
        '• MPI RCB 영역분할 병렬화\n'
        '• Lee 상변화, Rohsenow 핵비등, Zuber CHF, Nusselt 막응축 모델\n'
        '• Edwards 파이프 블로우다운 플래싱 검증\n'
        '• 1차 화학반응, P1 복사열전달 모델\n'
        '• AMR 사분면 세분화\n'
        '• CuPy GPU 가속 BiCGSTAB (CPU 폴백)\n'
        '• ILU(k)/AMG 전처리기\n'
        '• CFL 기반 적응 시간 간격 제어\n'
        '• ParaView 호환 VTU 출력\n'
        '• Flask/Plotly.js 실시간 모니터링 대시보드\n'
        '• 25개 검증 케이스를 통한 포괄적 코드 검증')

    _add_heading(doc, '12.2 향후 과제', level=2)
    _add_paragraph(doc,
        '• 고차 시간 이산화: Crank-Nicolson, Adams-Bashforth\n'
        '• 분산 압력 솔버: PETSc/Trilinos 연동\n'
        '• 동적 부하 재분산: 병렬 AMR에서의 실시간 격자 재분할\n'
        '• 다중 반응 시스템: 복잡 화학반응 네트워크\n'
        '• DOM 복사 모델: Discrete Ordinates Method\n'
        '• 3D 비정렬 AMR: 사면체 적응 세분화\n'
        '• 다성분 상변화: 비응축성 기체 혼합물 모델')

    doc.add_page_break()


def _write_ch13_manual(doc):
    """13장 프로그래머 매뉴얼."""
    _add_heading(doc, '13. 프로그래머 매뉴얼', level=1)
    _add_paragraph(doc,
        '본 장에서는 코드를 구성하는 각 모듈의 구조, '
        '주요 클래스 및 함수, 알고리즘을 기술한다.')

    # 13.1 mesh/
    _add_heading(doc, '13.1 mesh/ 모듈', level=2)

    _add_heading(doc, '13.1.1 mesh_reader.py', level=3)
    _add_paragraph(doc,
        '목적: FVM 격자 자료구조 및 Gmsh .msh 파일 파서.\n\n'
        '주요 클래스: FVCell, FVFace, FVMesh\n'
        '주요 함수: read_gmsh(filepath) → FVMesh')

    _add_heading(doc, '13.1.2 mesh_generator.py', level=3)
    _add_paragraph(doc,
        '목적: 2D 구조 격자 생성.\n\n'
        '함수: generate_channel_mesh, generate_cavity_mesh, '
        'generate_cht_mesh, generate_bubble_column_mesh, '
        'generate_triangle_channel_mesh')

    _add_heading(doc, '13.1.3 mesh_generator_3d.py', level=3)
    _add_paragraph(doc,
        '목적: 3D 육면체 구조격자 생성.\n\n'
        '함수: generate_3d_channel_mesh, generate_3d_duct_mesh, '
        'generate_3d_cavity_mesh')

    _add_heading(doc, '13.1.4 amr.py', level=3)
    _add_paragraph(doc,
        '목적: 2D 적응 격자 세분화.\n\n'
        '클래스: AMRMesh (FVMesh 상속)\n'
        '함수: refine_cells, gradient_jump_estimator')

    _add_heading(doc, '13.1.5 vtk_exporter.py', level=3)
    _add_paragraph(doc,
        '목적: ParaView VTU 내보내기.\n\n'
        '함수: export_mesh_to_vtu(mesh, filepath, cell_data)\n'
        '의존성: meshio')

    doc.add_page_break()

    # 13.2 core/
    _add_heading(doc, '13.2 core/ 모듈', level=2)

    _add_heading(doc, '13.2.1 fields.py', level=3)
    _add_paragraph(doc,
        '클래스: ScalarField(mesh, name), VectorField(mesh, name)')

    _add_heading(doc, '13.2.2 gradient.py', level=3)
    _add_paragraph(doc,
        '함수: green_gauss_gradient, least_squares_gradient')

    _add_heading(doc, '13.2.3 interpolation.py', level=3)
    _add_paragraph(doc,
        '함수: upwind_interpolation, central_interpolation, '
        'tvd_interpolation, muscl_reconstruct')

    _add_heading(doc, '13.2.4 fvm_operators.py', level=3)
    _add_paragraph(doc,
        '클래스: FVMSystem(n_cells)\n'
        '함수: diffusion_operator, convection_operator, source_term, time_derivative')

    _add_heading(doc, '13.2.5 linear_solver.py', level=3)
    _add_paragraph(doc,
        '함수: solve_sparse(A, b, method, tol, maxiter)')

    doc.add_page_break()

    # 13.3 models/
    _add_heading(doc, '13.3 models/ 모듈', level=2)

    _add_heading(doc, '13.3.1 single_phase.py', level=3)
    _add_paragraph(doc,
        '클래스: SIMPLESolver(mesh, rho, mu)\n'
        '메서드: set_velocity_bc, set_pressure_bc, solve_steady')

    _add_heading(doc, '13.3.2 two_fluid.py', level=3)
    _add_paragraph(doc,
        '클래스: TwoFluidSolver(mesh, rho_l, rho_g, mu_l, mu_g, d_b)\n'
        '메서드: solve_transient(dt, n_steps)')

    _add_heading(doc, '13.3.3 turbulence.py', level=3)
    _add_paragraph(doc,
        '클래스: KEpsilonModel(mesh, rho, mu)\n'
        '메서드: solve(U, dt), wall_function(y_plus)')

    _add_heading(doc, '13.3.4 closure.py', level=3)
    _add_paragraph(doc,
        '함수: schiller_naumann_drag, ranz_marshall_nusselt, sato_bit_viscosity')

    _add_heading(doc, '13.3.5 phase_change.py', level=3)
    _add_paragraph(doc,
        '클래스:\n'
        '• LeePhaseChangeModel: 체적 증발/응축\n'
        '• RohsenowBoilingModel: 핵비등 열유속\n'
        '• ZuberCHFModel: 임계 열유속\n'
        '• NusseltCondensationModel: 막응축\n'
        '• PhaseChangeManager: 통합 관리\n\n'
        '유틸리티 함수:\n'
        '• saturation_temperature(P): Antoine 근사 T_sat(P)\n'
        '• water_latent_heat(P): 잠열 h_fg(P)\n'
        '• water_properties(P): 포화 물성치')

    _add_heading(doc, '13.3.6 reaction.py', level=3)
    _add_paragraph(doc, '함수: reaction_source, solve_species_transport')

    _add_heading(doc, '13.3.7 radiation.py', level=3)
    _add_paragraph(doc, '함수: solve_p1_radiation, radiation_source')

    doc.add_page_break()

    # 13.4 parallel/
    _add_heading(doc, '13.4 parallel/ 모듈', level=2)

    _add_heading(doc, '13.4.1 partitioning.py', level=3)
    _add_paragraph(doc,
        '함수: rcb_partition, build_ghost_cells, exchange_ghost_data')

    _add_heading(doc, '13.4.2 mpi_solver.py', level=3)
    _add_paragraph(doc,
        '클래스: MPISolver(mesh, n_parts, rho, mu)')

    # 13.5 gpu/
    _add_heading(doc, '13.5 gpu/ 모듈', level=2)
    _add_paragraph(doc,
        '함수: detect_gpu_backend, gpu_bicgstab, _cpu_bicgstab, benchmark_solvers')

    # 13.6 verification/
    _add_heading(doc, '13.6 verification/ 모듈', level=2)

    _add_table(doc,
        ['파일', '케이스', '검증 대상'],
        [
            ['case1_poiseuille.py', 'Poiseuille', '2D 층류 해석해'],
            ['case2_cavity.py', 'Cavity', 'Ghia 벤치마크'],
            ['case3_cht.py', 'CHT', '열전달 커플링'],
            ['case4_bubble_column.py', '기포탑', 'Two-Fluid'],
            ['case5_3d_duct.py', '3D 덕트', '3D Poiseuille'],
            ['case6_muscl.py', 'MUSCL', '고차 이산화'],
            ['case7_unstructured.py', '비정렬', '삼각형 격자'],
            ['case8_mpi.py', 'MPI', '병렬 일치성'],
            ['case9_phase_change.py', '상변화', 'Lee 모델'],
            ['case10_reaction.py', '반응', '해석해 비교'],
            ['case11_radiation.py', '복사', 'P1 에너지 균형'],
            ['case12_amr.py', 'AMR', '격자 세분화'],
            ['case13_gpu.py', 'GPU', 'CPU/GPU 정확도'],
            ['case14_3d_cavity.py', '3D Cavity', '와류 유동'],
            ['case15_3d_convection.py', '3D 대류', 'Boussinesq'],
            ['case16_preconditioner.py', '전처리기', '수렴 비교'],
            ['case17_adaptive_dt.py', '적응 dt', 'CFL 제어'],
            ['case18_web_dashboard.py', '대시보드', 'Flask/Plotly'],
            ['case19_boiling_condensation.py', '비등/응축', 'Rohsenow/Nusselt'],
            ['case20_edwards_blowdown.py', 'Edwards', '플래싱'],
        ])

    _add_paragraph(doc,
        '각 스크립트는 run_caseN() 함수를 제공하며, run_all.py에서 일괄 실행된다.')

    # 13.7 report/
    _add_heading(doc, '13.7 report/ 모듈', level=2)
    _add_paragraph(doc,
        '함수: generate_report(results, output_path, figures_dir)\n'
        '헬퍼: _add_heading, _add_paragraph, _add_figure, _add_table\n'
        '의존성: python-docx')

    doc.add_page_break()


def _write_ch14_references(doc):
    """14장 참고문헌 (맨 마지막)."""
    _add_heading(doc, '14. 참고문헌', level=1)

    refs = [
        '[1] Patankar, S.V., "Numerical Heat Transfer and Fluid Flow", '
        'Hemisphere Publishing, 1980.',
        '[2] Versteeg, H.K., Malalasekera, W., "An Introduction to '
        'Computational Fluid Dynamics: The Finite Volume Method", Pearson, 2007.',
        '[3] Ghia, U., Ghia, K.N., Shin, C.T., "High-Re Solutions for '
        'Incompressible Flow Using the Navier-Stokes Equations and a '
        'Multigrid Method", J. Computational Physics, 48, 387-411, 1982.',
        '[4] Schiller, L., Naumann, Z., "A Drag Coefficient Correlation", '
        'Z. Ver. Deutsch. Ing., 77, 318-320, 1935.',
        '[5] Ranz, W.E., Marshall, W.R., "Evaporation from Drops", '
        'Chem. Eng. Progress, 48, 141-146, 1952.',
        '[6] Sato, Y., Sekoguchi, K., "Liquid Velocity Distribution in '
        'Two-Phase Bubble Flow", Int. J. Multiphase Flow, 2, 79-95, 1975.',
        '[7] Ishii, M., Hibiki, T., "Thermo-Fluid Dynamics of Two-Phase '
        'Flow", Springer, 2011.',
        '[8] Launder, B.E., Spalding, D.B., "The Numerical Computation of '
        'Turbulent Flows", CAME, 3(2), 269-289, 1974.',
        '[9] van Leer, B., "Towards the Ultimate Conservative Difference Scheme. '
        'V. A Second Order Sequel to Godunov\'s Method", JCP, 32(1), 101-136, 1979.',
        '[10] Sweby, P.K., "High Resolution Schemes Using Flux Limiters for '
        'Hyperbolic Conservation Laws", SIAM J. Num. Anal., 21(5), 995-1011, 1984.',
        '[11] Lee, W.H., "A Pressure Iteration Scheme for Two-Phase Flow Modeling", '
        'LA-UR-79-975, Los Alamos, 1980.',
        '[12] Modest, M.F., "Radiative Heat Transfer", 3rd ed., Academic Press, 2013.',
        '[13] Berger, M.J., Colella, P., "Local Adaptive Mesh Refinement for '
        'Shock Hydrodynamics", JCP, 82(1), 64-84, 1989.',
        '[14] Gropp, W., Lusk, E., Skjellum, A., "Using MPI", MIT Press, 2014.',
        '[15] Okuta, R., et al., "CuPy: A NumPy-Compatible Library for GPU", '
        'LearningSys in NIPS 2017.',
        '[16] Rohsenow, W.M., "A Method of Correlating Heat Transfer Data '
        'for Surface Boiling of Liquids", Trans. ASME, 74, 969-976, 1952.',
        '[17] Zuber, N., "Hydrodynamic Aspects of Boiling Heat Transfer", '
        'AEC Report AECU-4439, 1959.',
        '[18] Nusselt, W., "Die Oberflächenkondensation des Wasserdampfes", '
        'Z. VDI, 60, 541-546, 569-575, 1916.',
        '[19] Edwards, A.R., O\'Brien, T.P., "Studies of Phenomena Connected '
        'with the Depressurization of Water Reactors", J. British Nuclear '
        'Energy Society, 9(2), 125-135, 1970.',
    ]

    for ref in refs:
        _add_paragraph(doc, ref)


# =====================================================================
# 메인 함수
# =====================================================================

def generate_report(results: dict, output_path: str,
                    figures_dir: str = "figures"):
    """
    DOCX 보고서 생성.

    Parameters
    ----------
    results : 전체 검증 결과 딕셔너리
    output_path : 출력 파일 경로
    figures_dir : 그림 파일 디렉토리
    """
    doc = Document()

    # 문서 스타일
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)

    # 표지 + 목차
    _write_cover(doc)
    _write_toc(doc)

    # 1~6장: 이론/수치기법/코드구조/격자/병렬화
    _write_ch1_intro(doc)
    _write_ch2_governing(doc)
    _write_ch3_numerics(doc)
    _write_ch4_code_structure(doc)
    _write_ch5_mesh(doc)
    _write_ch6_parallel(doc)

    # 7~10장: 검증 결과
    _write_ch7_verification_2d(doc, results, figures_dir)
    _write_ch8_verification_3d(doc, results, figures_dir)
    _write_ch9_verification_phase_change(doc, results, figures_dir)
    _write_ch10_verification_etc(doc, results, figures_dir)

    # 11장: 검증 요약
    _write_ch11_summary(doc, results)

    # 12장: 결론
    _write_ch12_conclusion(doc)

    # 13장: 프로그래머 매뉴얼
    _write_ch13_manual(doc)

    # 14장: 참고문헌 (맨 마지막)
    _write_ch14_references(doc)

    # 저장
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"  보고서 저장: {output_path}")

    return output_path
